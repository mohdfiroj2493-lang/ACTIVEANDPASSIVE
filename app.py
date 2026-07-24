import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import date
import os

st.set_page_config(
    page_title="Lateral Earth Pressure and Deflection Calculator - English Units",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
.main-header{background:linear-gradient(135deg,#1a3a5c 0%,#2d6a9f 100%);padding:20px 30px;border-radius:12px;margin-bottom:24px;color:white}.main-header h1{margin:0;font-size:2rem}.main-header p{margin:4px 0 0;opacity:.85;font-size:.95rem}.result-card{background:#f0f7ff;border-left:5px solid #2d6a9f;border-radius:8px;padding:16px 20px;margin:8px 0}.result-card h4{margin:0 0 8px;color:#1a3a5c}.method-badge{display:inline-block;padding:3px 10px;border-radius:20px;font-size:.78rem;font-weight:600;margin-bottom:6px}.rankine{background:#dbeafe;color:#1e40af}.coulomb{background:#dcfce7;color:#166534}.atrest{background:#fef9c3;color:#854d0e}.note{background:#fff7ed;border:1px solid #fed7aa;border-radius:8px;padding:10px 14px;color:#7c2d12}.stTabs [data-baseweb="tab-list"]{gap:8px}.stTabs [data-baseweb="tab"]{background:#f1f5f9;border-radius:8px 8px 0 0;padding:8px 20px}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
  <h1>🏗️ Lateral Earth Pressure and Deflection Calculator</h1>
  <p>English Units · Multi-Layer Soil · Water Table · Rankine · Coulomb · At-Rest · Surcharge · Elastic Deflection</p>
</div>
""", unsafe_allow_html=True)

# -----------------------------
# Soil layer table defaults
# -----------------------------
def default_layers(n, height):
    bottoms = np.linspace(height / n, height, n)
    return pd.DataFrame({
        "Bottom Depth (ft)": np.round(bottoms, 2),
        "Moist Unit Weight γm (pcf)": [120.0] * n,
        "Saturated Unit Weight γsat (pcf)": [125.0] * n,
        "Friction Angle φ (deg)": [30.0] * n,
        "Cohesion c (psf)": [0.0] * n,
    })

# -----------------------------
# Sidebar inputs
# -----------------------------
with st.sidebar:
    with st.form("analysis_parameters_form", clear_on_submit=False):
        st.header("⚙️ Input Parameters")

        st.subheader("🧱 Wall Geometry")
        H = st.number_input("Wall Height H (ft)", 3.0, 100.0, 20.0, 1.0)
        alpha_deg = st.number_input("Wall Inclination α (°)", -20.0, 20.0, 0.0, 1.0,
                                    help="Angle of wall face from vertical (+ = leaning into backfill)")
        beta_deg = st.number_input("Backfill Slope β (°)", 0.0, 30.0, 0.0, 1.0)
        delta_deg = st.number_input("Wall Friction δ (°), for Coulomb", 0.0, 35.0, 15.0, 1.0)

        st.subheader("🟫 Passive Pressure Zone")
        passive_start_depth = st.number_input(
            "Passive Pressure Start Depth from Top (ft)",
            min_value=0.0,
            max_value=float(H),
            value=0.0,
            step=0.5,
            format="%.2f",
            help="Passive resistance is set to zero above this depth. Below this depth, passive pressure is calculated using depth measured from this point."
        )

        st.subheader("💧 Groundwater")
        include_water = st.checkbox("Include water table", True)
        water_table = st.number_input("Water Table Depth from Top (ft)", 0.0, float(H), min(10.0, float(H)), 0.5)
        gamma_w = st.number_input("Unit Weight of Water γw (pcf)", 60.0, 64.0, 62.4, 0.1)

        st.subheader("🪨 Soil Properties / Layers")
        st.caption("Use Add Layer to create more layers. Enter each layer bottom depth from the top; each layer thickness can be any value greater than 0 ft.")

        if "soil_layers" not in st.session_state:
            st.session_state.soil_layers = default_layers(1, float(H)).to_dict("records")

        if st.form_submit_button("➕ Add Layer", use_container_width=True):
            st.session_state["analysis_has_run"] = False
            layers = st.session_state.soil_layers
            if len(layers) < 8:
                if layers:
                    previous_top = 0.0 if len(layers) == 1 else float(layers[-2]["Bottom Depth (ft)"])
                    split_depth = round((previous_top + float(H)) / 2.0, 2)
                    layers[-1]["Bottom Depth (ft)"] = max(previous_top + 0.1, min(split_depth, float(H) - 0.1))
                    template = layers[-1].copy()
                    template["Bottom Depth (ft)"] = float(H)
                else:
                    template = default_layers(1, float(H)).iloc[0].to_dict()
                layers.append(template)
                st.session_state.soil_layers = layers
            else:
                st.warning("Maximum 8 layers allowed.")

        if len(st.session_state.soil_layers) > 1:
            if st.form_submit_button("➖ Remove Last Layer", use_container_width=True):
                st.session_state.soil_layers = st.session_state.soil_layers[:-1]
                st.session_state.soil_layers[-1]["Bottom Depth (ft)"] = float(H)
                st.session_state["analysis_has_run"] = False

        updated_layers = []
        previous_bottom = 0.0
        for i, layer in enumerate(st.session_state.soil_layers):
            with st.container(border=True):
                st.markdown(f"**Layer {i + 1}**")
                # Bottom depth is an editable layer boundary.
                # Each layer only needs a thickness greater than zero; the last
                # layer is no longer locked to the wall height. If the final
                # boundary is above the wall base, the calculation extends the
                # last entered layer properties down to H.
                bottom_min = min(float(H), previous_bottom + 0.01)
                bottom_max = float(H)
                bottom_default = float(layer.get("Bottom Depth (ft)", bottom_max))
                bottom_default = min(max(bottom_default, bottom_min), bottom_max)

                bottom_value = st.number_input(
                    "Bottom Depth from Top (ft)",
                    min_value=bottom_min,
                    max_value=bottom_max,
                    value=bottom_default,
                    step=0.1,
                    format="%.2f",
                    key=f"layer_{i}_bottom",
                    help="Enter the bottom depth of this layer. Layer thickness must be greater than 0 ft."
                )

                gamma_m_value = st.number_input(
                    "Moist Unit Weight γm (pcf)",
                    min_value=50.0,
                    max_value=160.0,
                    value=float(layer.get("Moist Unit Weight γm (pcf)", 120.0)),
                    step=1.0,
                    key=f"layer_{i}_gamma_m",
                )
                gamma_sat_value = st.number_input(
                    "Saturated Unit Weight γsat (pcf)",
                    min_value=50.0,
                    max_value=170.0,
                    value=float(layer.get("Saturated Unit Weight γsat (pcf)", 125.0)),
                    step=1.0,
                    key=f"layer_{i}_gamma_sat",
                )
                phi_value = st.number_input(
                    "Friction Angle φ (deg)",
                    min_value=0.0,
                    max_value=50.0,
                    value=float(layer.get("Friction Angle φ (deg)", 30.0)),
                    step=1.0,
                    key=f"layer_{i}_phi",
                )
                c_value = st.number_input(
                    "Cohesion c (psf)",
                    min_value=0.0,
                    max_value=5000.0,
                    value=float(layer.get("Cohesion c (psf)", 0.0)),
                    step=50.0,
                    key=f"layer_{i}_c",
                )

            previous_bottom = float(bottom_value)
            updated_layers.append({
                "Bottom Depth (ft)": float(bottom_value),
                "Moist Unit Weight γm (pcf)": float(gamma_m_value),
                "Saturated Unit Weight γsat (pcf)": float(gamma_sat_value),
                "Friction Angle φ (deg)": float(phi_value),
                "Cohesion c (psf)": float(c_value),
            })

        st.session_state.soil_layers = updated_layers
        layer_df = pd.DataFrame(updated_layers)

        st.subheader("📦 Surcharge")
        surcharge_options = {
            "None": "None",
            "Uniform (Kq)": "Uniform",
            "Line Load (NAVFAC/Boussinesq)": "Line Load",
            "Point Load (NAVFAC/Boussinesq)": "Point Load",
            "Strip Load (FHWA/WALLPRES)": "Strip Load",
            "AASHTO Strip/Isolated Footing (2:1 Distribution)": "AASHTO",
        }
        surcharge_label = st.selectbox("Surcharge Type", list(surcharge_options.keys()))
        surcharge_type = surcharge_options[surcharge_label]

        q_uniform = 0.0
        Q_line = 0.0
        x_line = 5.0
        Q_point = 0.0
        x_point = 5.0
        q_strip = 0.0
        x_strip_near = 2.0
        strip_width = 10.0
        aashto_load_type = "Strip footing"
        aashto_Pv = 0.0
        aashto_bf = 2.0
        aashto_L = 10.0
        aashto_d = 5.0

        if surcharge_type == "Uniform":
            q_uniform = st.number_input("Uniform Surcharge q (psf)", 0.0, 10000.0, 250.0, 25.0)
            st.caption("Calculated separately as Δp = K × q.")
        elif surcharge_type == "Line Load":
            Q_line = st.number_input("Line Load Q (lb/ft)", 0.0, 50000.0, 2000.0, 100.0)
            x_line = st.number_input("Distance from Wall x (ft)", 0.5, 100.0, 6.0, 0.5)
            st.caption("Calculated separately using the NAVFAC/Boussinesq wall-pressure equation.")
        elif surcharge_type == "Point Load":
            Q_point = st.number_input("Point Load Q (lb)", 0.0, 200000.0, 10000.0, 500.0)
            x_point = st.number_input("Distance from Wall x (ft)", 0.5, 100.0, 6.0, 0.5)
            st.caption("Calculated separately using the NAVFAC/Boussinesq wall-pressure equation.")
        elif surcharge_type == "Strip Load":
            q_strip = st.number_input("Strip Load q (psf)", 0.0, 10000.0, 600.0, 25.0)
            x_strip_near = st.number_input("Distance to Near Edge x1 (ft)", 0.0, 200.0, 2.0, 0.5)
            strip_width = st.number_input("Strip Width B (ft)", 0.1, 500.0, 30.0, 0.5)
            st.caption("Calculated separately with the WALLPRES strip-load approach: x2 = x1 + B.")
        elif surcharge_type == "AASHTO":
            aashto_load_type = st.selectbox(
                "AASHTO Load Type",
                ["Strip footing", "Isolated rectangular footing", "Point load"],
            )
            aashto_d = st.number_input(
                "Distance d from wall back face to load centroid (ft)",
                0.0, 500.0, 5.0, 0.5,
            )
            if aashto_load_type == "Strip footing":
                aashto_Pv = st.number_input("Strip load Pv (lb/ft)", 0.0, 500000.0, 2000.0, 100.0)
                aashto_bf = st.number_input("Loaded width bf (ft)", 0.0, 200.0, 4.0, 0.5)
                aashto_L = 0.0
                st.caption("AASHTO 2:1 distribution: Δσv = Pv / D1, then lateral surcharge Δp = K × Δσv.")
            elif aashto_load_type == "Isolated rectangular footing":
                aashto_Pv = st.number_input("Total vertical load P'v (lb)", 0.0, 5000000.0, 50000.0, 1000.0)
                aashto_bf = st.number_input("Footing width bf (ft)", 0.0, 200.0, 4.0, 0.5)
                aashto_L = st.number_input("Footing length L (ft)", 0.1, 500.0, 10.0, 0.5)
                st.caption("AASHTO 2:1 distribution: Δσv = P'v / [D1(L+z)], then lateral surcharge Δp = K × Δσv.")
            elif aashto_load_type == "Point load":
                aashto_Pv = st.number_input("Concentrated vertical load P'v (lb)", 0.0, 5000000.0, 50000.0, 1000.0)
                aashto_bf = 0.0
                aashto_L = 0.0
                st.caption("AASHTO 2:1 distribution for point load: use bf = 0 and Δσv = P'v / D1², then lateral surcharge Δp = K × Δσv.")
            st.caption("AASHTO surcharge pressure is set to zero above z2. Below z2: z2 = 2d - bf and D1 = (bf + z)/2 + d. For point load, bf = 0.")

        st.subheader("📐 Pile / Wall Flexural Properties")
        youngs_modulus_ksi = st.number_input(
            "Young's Modulus E (ksi)",
            min_value=1.0,
            max_value=100000.0,
            value=29000.0,
            step=100.0,
            format="%.1f",
            help="Elastic modulus of the pile or wall section. Steel is commonly entered as approximately 29,000 ksi.",
        )
        moment_of_inertia_in4 = st.number_input(
            "Moment of Inertia I (in⁴)",
            min_value=0.01,
            max_value=1.0e9,
            value=5000.0,
            step=100.0,
            format="%.2f",
            help="Gross or effective flexural moment of inertia of one pile or the selected wall strip.",
        )
        pile_tributary_width_ft = st.number_input(
            "Pile Tributary Width / Spacing (ft)",
            min_value=0.01,
            max_value=100.0,
            value=1.0,
            step=0.5,
            format="%.2f",
            help="Converts wall pressure in psf to the line load carried by one pile. Use pile spacing for a discrete pile wall; use 1.0 ft for a one-foot wall strip.",
        )

        deflection_fixity_mode = st.selectbox(
            "Deflection Boundary Condition",
            [
                "Point of fixity within embedment (Shoring Suite style)",
                "Fixed at pile toe",
            ],
            index=0,
            help=(
                "The point-of-fixity option imposes zero rotation and zero deflection at a selected point within "
                "the embedded length. The toe-fixed option imposes those conditions at the pile toe."
            ),
        )

        embedment_depth_ft = max(float(H) - float(passive_start_depth), 0.0)
        if deflection_fixity_mode.startswith("Point of fixity"):
            fixity_percent = st.number_input(
                "Point of Fixity below Excavation (% of Embedment)",
                min_value=0.0,
                max_value=100.0,
                value=60.0,
                step=1.0,
                format="%.1f",
                help=(
                    "Measured downward from the passive-pressure start/excavation depth. "
                    "A value between about 50% and 67% is commonly used for comparison with Shoring Suite."
                ),
            )
            fixity_fraction = float(fixity_percent) / 100.0
            point_of_fixity_depth = float(passive_start_depth) + fixity_fraction * embedment_depth_ft
            deflection_boundary_label = (
                f"Virtual fixity at {point_of_fixity_depth:.2f} ft "
                f"({fixity_percent:.1f}% of embedment; restrained below)"
            )
            st.caption(
                f"Embedment = {embedment_depth_ft:.2f} ft; calculated point of fixity = "
                f"{point_of_fixity_depth:.2f} ft below the wall top."
            )
        else:
            fixity_percent = 100.0
            fixity_fraction = 1.0
            point_of_fixity_depth = float(H)
            deflection_boundary_label = f"Pile toe fixed at {point_of_fixity_depth:.2f} ft"

        st.caption(
            "Deflection is calculated by double integration of M/EI above the selected virtual fixity. "
            "The fixity point has zero rotation and zero deflection. For the Shoring Suite-style option, "
            "the embedded segment below fixity is treated as restrained and is plotted at zero deflection. "
            "The imposed passive-pressure diagram is not iterated with wall movement."
        )

        st.subheader("🏛️ USACE Cantilever Method")
        run_usace_analysis = st.checkbox(
            "Run USACE cantilever equilibrium analysis",
            value=True,
            help=(
                "Implements the cantilever-wall stability procedure in USACE EM 1110-2-2504. "
                "The solver determines the required penetration and transition point from simultaneous "
                "horizontal-force and moment equilibrium."
            ),
        )
        usace_passive_fs = st.number_input(
            "USACE Passive-Pressure Safety Factor FSp",
            min_value=1.0,
            max_value=5.0,
            value=1.50,
            step=0.05,
            format="%.2f",
            help=(
                "USACE reduces passive soil strength using tan(phi_eff)=tan(phi)/FSp and c_eff=c/FSp. "
                "For a usual-condition retaining wall in free-draining soil, EM 1110-2-2504 lists 1.50."
            ),
        )
        usace_default_max_embedment = max(
            20.0,
            3.0 * max(float(H) - float(passive_start_depth), 1.0),
            3.0 * max(float(passive_start_depth), 1.0),
        )
        usace_max_embedment_ft = st.number_input(
            "USACE Maximum Trial Embedment (ft)",
            min_value=1.0,
            max_value=500.0,
            value=float(min(usace_default_max_embedment, 500.0)),
            step=1.0,
            format="%.1f",
            help=(
                "Upper search bound for the required penetration below the excavation/passive-pressure start depth. "
                "The last entered soil layer is extended below the entered wall bottom for this trial calculation."
            ),
        )
        usace_solver_points = st.slider(
            "USACE Final Diagram Points",
            min_value=201,
            max_value=1201,
            value=301,
            step=100,
        )
        st.caption(
            "The USACE transition point is near the point of zero displacement. For the optional elastic-deflection "
            "diagram, it is used as an idealized zero-rotation and zero-deflection point; that structural boundary "
            "condition is an approximation beyond the rotational-stability procedure in the manual. Passive pressure "
            "continues to use this app's Kp = 1/Ka convention after the USACE strength reduction is applied."
        )

        st.subheader("📊 Display Options")
        n_points = st.slider("Pressure Diagram Points", 50, 500, 151)
        show_cohesion = st.checkbox("Include cohesion term", True)
        show_total_pressure = st.checkbox("Show total pressure including water", True)


        st.markdown("---")
        run_analysis_button = st.form_submit_button(
            "▶️ Run Analysis",
            type="primary",
            use_container_width=True,
        )
        st.caption(
            "Change any parameters above, then click Run Analysis. "
            "Inputs inside this form do not rerun the calculation until submitted."
        )

        if run_analysis_button:
            st.session_state["analysis_has_run"] = True
            st.session_state["analysis_run_number"] = (
                int(st.session_state.get("analysis_run_number", 0)) + 1
            )


# Do not execute the engineering calculations until the user submits the form.
if not st.session_state.get("analysis_has_run", False):
    st.info(
        "Enter or revise the parameters in the sidebar, then click **Run Analysis**. "
        "The pressure, USACE, shear, moment, and deflection calculations will run only after submission."
    )
    st.stop()

# -----------------------------
# Soil layer cleanup
# -----------------------------
layer_df = layer_df.sort_values("Bottom Depth (ft)").reset_index(drop=True)
if layer_df.iloc[-1]["Bottom Depth (ft)"] < H:
    st.warning("The last soil layer bottom is less than the wall height. The app will extend the last layer to the wall base.")
    layer_df.loc[layer_df.index[-1], "Bottom Depth (ft)"] = H
elif layer_df.iloc[-1]["Bottom Depth (ft)"] > H:
    layer_df.loc[layer_df.index[-1], "Bottom Depth (ft)"] = H

# -----------------------------
# Calculations
# -----------------------------
alpha = np.radians(alpha_deg)
beta = np.radians(beta_deg)
delta = np.radians(delta_deg)
depths = np.linspace(0.0, H, int(n_points))


def safe_sqrt(x):
    return np.sqrt(np.maximum(x, 0.0))


def rankine_Ka(phi, beta):
    if abs(beta) < 1e-9:
        return np.tan(np.pi / 4 - phi / 2) ** 2
    rad = np.cos(beta) ** 2 - np.cos(phi) ** 2
    if rad < 0:
        return np.nan
    num = np.cos(beta) - np.sqrt(rad)
    den = np.cos(beta) + np.sqrt(rad)
    return np.cos(beta) * num / den


def rankine_Kp(phi, beta):
    if abs(beta) < 1e-9:
        return np.tan(np.pi / 4 + phi / 2) ** 2
    rad = np.cos(beta) ** 2 - np.cos(phi) ** 2
    if rad < 0:
        return np.nan
    num = np.cos(beta) + np.sqrt(rad)
    den = np.cos(beta) - np.sqrt(rad)
    return np.cos(beta) * num / den


def coulomb_Ka(phi, delta, alpha, beta):
    try:
        rad = (np.sin(phi + delta) * np.sin(phi - beta)) / (np.cos(delta + alpha) * np.cos(alpha - beta))
        if rad < 0:
            return np.nan
        num = np.cos(phi - alpha) ** 2
        den = np.cos(alpha) ** 2 * np.cos(delta + alpha) * (1 + np.sqrt(rad)) ** 2
        return num / den
    except Exception:
        return np.nan


def coulomb_Kp(phi, delta, alpha, beta):
    try:
        rad = (np.sin(phi + delta) * np.sin(phi + beta)) / (np.cos(delta - alpha) * np.cos(alpha - beta))
        if rad < 0:
            return np.nan
        num = np.cos(phi + alpha) ** 2
        den = np.cos(alpha) ** 2 * np.cos(delta - alpha) * (1 - np.sqrt(rad)) ** 2
        if abs(den) < 1e-12:
            return np.nan
        return num / den
    except Exception:
        return np.nan


def K0(phi, beta):
    return (1 - np.sin(phi)) * (1 + np.sin(beta))


def layer_index_at_depth(z):
    bottoms = layer_df["Bottom Depth (ft)"].to_numpy(dtype=float)
    return min(int(np.searchsorted(bottoms, z, side="left")), len(bottoms) - 1)


def properties_at_depth(z):
    row = layer_df.iloc[layer_index_at_depth(z)]
    phi = np.radians(float(row["Friction Angle φ (deg)"]))
    c = float(row["Cohesion c (psf)"])
    return phi, c


def vertical_stresses_at_depth(z):
    """Return total vertical stress, pore pressure, and effective vertical stress at depth z.

    For the standard wall calculation, z is no greater than H. The USACE
    penetration solver may trial a deeper wall bottom; in that case the last
    entered soil layer is extended below its entered bottom depth.
    """
    if z <= 0:
        return 0.0, 0.0, 0.0

    def segment_total_stress(seg_top, seg_bot, gamma_m, gamma_sat):
        if seg_bot <= seg_top:
            return 0.0
        if not include_water:
            return gamma_m * (seg_bot - seg_top)
        if seg_bot <= water_table:
            return gamma_m * (seg_bot - seg_top)
        if seg_top >= water_table:
            return gamma_sat * (seg_bot - seg_top)
        return (
            gamma_m * (water_table - seg_top)
            + gamma_sat * (seg_bot - water_table)
        )

    total = 0.0
    top = 0.0
    last_row = layer_df.iloc[-1]

    for _, row in layer_df.iterrows():
        bottom = float(row["Bottom Depth (ft)"])
        if z <= top:
            break
        seg_bot = min(float(z), bottom)
        if seg_bot > top:
            total += segment_total_stress(
                top,
                seg_bot,
                float(row["Moist Unit Weight γm (pcf)"]),
                float(row["Saturated Unit Weight γsat (pcf)"]),
            )
        top = bottom
        last_row = row
        if top >= z:
            break

    # Extend the last entered layer when the USACE solver trials a bottom below H.
    if top < z:
        total += segment_total_stress(
            top,
            float(z),
            float(last_row["Moist Unit Weight γm (pcf)"]),
            float(last_row["Saturated Unit Weight γsat (pcf)"]),
        )

    u = gamma_w * max(0.0, float(z) - float(water_table)) if include_water else 0.0
    effective = max(total - u, 0.0)
    return total, u, effective


def vertical_stresses_between_depths(z_top, z_bot):
    """Returns total, pore pressure, and effective vertical stress for the soil column between z_top and z_bot.

    This is used for passive pressure when passive resistance starts below the wall top,
    such as at an excavation subgrade/dredge line.
    """
    if z_bot <= z_top:
        return 0.0, 0.0, 0.0
    total_bot, _, _ = vertical_stresses_at_depth(z_bot)
    total_top, _, _ = vertical_stresses_at_depth(z_top)
    total_increment = max(total_bot - total_top, 0.0)

    if include_water:
        effective_water_table = max(float(water_table), float(z_top))
        u_increment = gamma_w * max(0.0, z_bot - effective_water_table)
    else:
        u_increment = 0.0

    effective_increment = max(total_increment - u_increment, 0.0)
    return total_increment, u_increment, effective_increment


def surcharge_pressure(z_arr, stype, q=0, Q=0, x=1, K_arr=None, x1=0.0, width=0.0,
                       aashto_load_type="Strip footing", aashto_Pv=0.0, aashto_bf=0.0,
                       aashto_L=0.0, aashto_d=0.0):
    """Returns surcharge-only lateral wall pressure in psf.

    Uniform surcharge uses Δp = Kq. Line, point, and strip load use the
    FHWA/WALLPRES-style equations and are kept separate from basic earth
    pressure. For strip load: x1 = distance to near edge, x2 = x1 + width.

    AASHTO option uses the 2:1 effective-width method from the user-provided
    sketch for strip footings, isolated rectangular footings, and concentrated
    point loads. The surcharge pressure is set to zero above z2. Below z2:
    z2 = 2d - bf and D1 = (bf + z)/2 + d. For point loads, bf = 0 and
    Δσv = P'v / D1². Vertical stress is converted to lateral surcharge as
    Δp = K × Δσv.
    """
    p = np.zeros_like(z_arr, dtype=float)
    if K_arr is None:
        K_arr = np.ones_like(z_arr, dtype=float)
    if stype == "Uniform":
        p = K_arr * q
    elif stype == "AASHTO":
        bf = max(float(aashto_bf), 0.0)
        d = max(float(aashto_d), 0.0)
        L = max(float(aashto_L), 0.0)
        Pv = max(float(aashto_Pv), 0.0)
        z2 = max(0.0, 2.0 * d - bf)
        for i, z in enumerate(z_arr):
            # Per the requested AASHTO implementation, surcharge pressure is
            # zero above z2 and is calculated only below z2.
            if Pv <= 0 or z <= z2:
                continue
            D1 = (bf + z) / 2.0 + d
            D1 = max(D1, 1e-6)
            if aashto_load_type == "Strip footing":
                delta_sigma_v = Pv / D1
            elif aashto_load_type == "Isolated rectangular footing":
                delta_sigma_v = Pv / (D1 * max(L + z, 1e-6))
            elif aashto_load_type == "Point load":
                if z <= 0:
                    continue
                delta_sigma_v = Pv / (D1 ** 2)
            else:
                delta_sigma_v = 0.0
            p[i] = K_arr[i] * delta_sigma_v
    elif stype == "Line Load":
        for i, z in enumerate(z_arr):
            if z <= 0:
                continue
            m = x / H
            n = z / H
            if m > 0.4:
                p[i] = (2 * Q / np.pi / H) * (m ** 2 * n / (m ** 2 + n ** 2) ** 2)
            else:
                p[i] = (0.203 * Q / H) * (n / (0.16 + n ** 2) ** 2)
    elif stype == "Point Load":
        for i, z in enumerate(z_arr):
            if z <= 0:
                continue
            m = x / H
            n = z / H
            if m > 0.4:
                p[i] = (3 * Q / (2 * np.pi * H ** 2)) * (m ** 2 * n ** 3 / (m ** 2 + n ** 2) ** 2.5)
            else:
                p[i] = (0.28 * Q / H ** 2) * (n ** 3 / (0.16 + n ** 2) ** 3)
    elif stype == "Strip Load":
        x2 = x1 + width
        for i, z in enumerate(z_arr):
            if z <= 0 or q <= 0 or width <= 0:
                continue
            # FHWA/WALLPRES strip surcharge equation for a rigid wall.
            # Angles are in radians here. The uploaded spreadsheet uses degrees and then converts back.
            beta_strip = np.arctan2(x2, z) - np.arctan2(x1, z)
            alpha_strip = np.arctan2(x1, z) + beta_strip / 2.0
            p[i] = (2.0 * q / np.pi) * (beta_strip - np.sin(beta_strip) * np.cos(2.0 * alpha_strip))
    return np.clip(p, 0.0, None)

sigma_v_total = np.zeros_like(depths)
u_water = np.zeros_like(depths)
sigma_v_eff = np.zeros_like(depths)
sigma_v_total_passive = np.zeros_like(depths)
u_water_passive = np.zeros_like(depths)
sigma_v_eff_passive = np.zeros_like(depths)
phi_arr = np.zeros_like(depths)
c_arr = np.zeros_like(depths)
Ka_r = np.zeros_like(depths)
Kp_r = np.zeros_like(depths)
Ka_c = np.zeros_like(depths)
Kp_c = np.zeros_like(depths)
K0_arr = np.zeros_like(depths)

for i, z in enumerate(depths):
    sigma_v_total[i], u_water[i], sigma_v_eff[i] = vertical_stresses_at_depth(z)
    sigma_v_total_passive[i], u_water_passive[i], sigma_v_eff_passive[i] = vertical_stresses_between_depths(passive_start_depth, z)
    phi_i, c_i = properties_at_depth(z)
    phi_arr[i] = phi_i
    c_arr[i] = c_i
    Ka_r[i] = rankine_Ka(phi_i, beta)
    Ka_c[i] = coulomb_Ka(phi_i, delta, alpha, beta)
    K0_arr[i] = K0(phi_i, beta)

# Replace invalid Coulomb active values with Rankine fallback so the app does not crash.
Ka_c = np.where(np.isfinite(Ka_c), Ka_c, Ka_r)

# Passive pressure in this app is calculated using the inverse of the active coefficient.
# This keeps the relationship Kp = 1 / Ka, so when c = 0 and no water is present:
# passive pressure = vertical effective stress / Ka.
Kp_r = np.where(Ka_r > 0, 1.0 / Ka_r, np.nan)
Kp_c = np.where(Ka_c > 0, 1.0 / Ka_c, Kp_r)

# Surcharge is calculated separately from earth/water pressure.
selected_Q = Q_line if surcharge_type == "Line Load" else Q_point
selected_x = x_line if surcharge_type == "Line Load" else x_point
sur_r = surcharge_pressure(depths, surcharge_type, q=q_uniform if surcharge_type == "Uniform" else q_strip,
                           Q=selected_Q, x=selected_x, K_arr=Ka_r,
                           x1=x_strip_near, width=strip_width,
                           aashto_load_type=aashto_load_type, aashto_Pv=aashto_Pv,
                           aashto_bf=aashto_bf, aashto_L=aashto_L, aashto_d=aashto_d)
sur_c = surcharge_pressure(depths, surcharge_type, q=q_uniform if surcharge_type == "Uniform" else q_strip,
                           Q=selected_Q, x=selected_x, K_arr=Ka_c,
                           x1=x_strip_near, width=strip_width,
                           aashto_load_type=aashto_load_type, aashto_Pv=aashto_Pv,
                           aashto_bf=aashto_bf, aashto_L=aashto_L, aashto_d=aashto_d)
sur_0 = surcharge_pressure(depths, surcharge_type, q=q_uniform if surcharge_type == "Uniform" else q_strip,
                           Q=selected_Q, x=selected_x, K_arr=K0_arr,
                           x1=x_strip_near, width=strip_width,
                           aashto_load_type=aashto_load_type, aashto_Pv=aashto_Pv,
                           aashto_bf=aashto_bf, aashto_L=aashto_L, aashto_d=aashto_d)

cohesion_term_r = 2 * c_arr * safe_sqrt(Ka_r) if show_cohesion else 0.0
cohesion_term_c = 2 * c_arr * safe_sqrt(Ka_c) if show_cohesion else 0.0

pa_rankine_eff = Ka_r * sigma_v_eff - cohesion_term_r
pa_coulomb_eff = Ka_c * sigma_v_eff - cohesion_term_c
pa_atrest_eff = K0_arr * sigma_v_eff

pa_rankine_eff_net = np.clip(pa_rankine_eff, 0, None)
pa_coulomb_eff_net = np.clip(pa_coulomb_eff, 0, None)
pa_atrest_eff_net = np.clip(pa_atrest_eff, 0, None)

passive_mask = depths >= passive_start_depth

# Passive effective soil pressure is calculated explicitly as sigma'_v / Ka.
# Because Kp = 1 / Ka, this is equivalent to Kp * sigma'_v for cohesionless soil.
# The effective vertical stress used here begins at the user-defined passive start depth.
passive_cohesion_r = (
    2.0 * c_arr * safe_sqrt(Kp_r)
    if show_cohesion
    else np.zeros_like(depths)
)
passive_cohesion_c = (
    2.0 * c_arr * safe_sqrt(Kp_c)
    if show_cohesion
    else np.zeros_like(depths)
)

pp_rankine_eff = np.where(
    passive_mask,
    np.divide(
        sigma_v_eff_passive,
        Ka_r,
        out=np.zeros_like(sigma_v_eff_passive),
        where=Ka_r > 0,
    ) + passive_cohesion_r,
    0.0,
)
pp_coulomb_eff = np.where(
    passive_mask,
    np.divide(
        sigma_v_eff_passive,
        Ka_c,
        out=np.zeros_like(sigma_v_eff_passive),
        where=Ka_c > 0,
    ) + passive_cohesion_c,
    0.0,
)

# Water pressure is kept separate from effective soil pressure.
# When total pressure is requested:
#     passive total pressure = passive effective soil pressure + pore-water pressure.
water_component = u_water if show_total_pressure else np.zeros_like(depths)
passive_water_component = u_water_passive if show_total_pressure else np.zeros_like(depths)
pa_rankine = pa_rankine_eff_net + water_component
pa_coulomb = pa_coulomb_eff_net + water_component
pa_atrest = pa_atrest_eff_net + water_component
pp_rankine = pp_rankine_eff + passive_water_component
pp_coulomb = pp_coulomb_eff + passive_water_component

# Optional combined diagrams/results for checking total demand.
pa_rankine_plus_surcharge = pa_rankine + sur_r
pa_coulomb_plus_surcharge = pa_coulomb + sur_c
pa_atrest_plus_surcharge = pa_atrest + sur_0


def resultant(p, z):
    F_lb_per_ft = float(np.trapezoid(p, z))
    if F_lb_per_ft <= 0:
        return 0.0, H / 3
    moment_about_base = float(np.trapezoid(p * (H - z), z))
    h_bar = moment_about_base / F_lb_per_ft
    return F_lb_per_ft / 1000.0, h_bar

Fa_r, ha_r = resultant(pa_rankine, depths)
Fa_c, ha_c = resultant(pa_coulomb, depths)
Fa_0, ha_0 = resultant(pa_atrest, depths)
Fs_r, hs_r = resultant(sur_r, depths)
Fs_c, hs_c = resultant(sur_c, depths)
Fs_0, hs_0 = resultant(sur_0, depths)
Ft_r, ht_r = resultant(pa_rankine_plus_surcharge, depths)
Ft_c, ht_c = resultant(pa_coulomb_plus_surcharge, depths)
Ft_0, ht_0 = resultant(pa_atrest_plus_surcharge, depths)
Fp_r, hp_r = resultant(pp_rankine, depths)
Fp_c, hp_c = resultant(pp_coulomb, depths)

# Net pressure for wall structural diagrams. Positive = active/surcharge side load;
# negative = passive resistance exceeds active + surcharge at that depth.
net_pressure_rankine = pa_rankine_plus_surcharge - pp_rankine
net_pressure_coulomb = pa_coulomb_plus_surcharge - pp_coulomb

def cumulative_trapezoid_np(y, x):
    """Cumulative trapezoidal integral from the first coordinate to each point."""
    y = np.asarray(y, dtype=float)
    x = np.asarray(x, dtype=float)
    out = np.zeros_like(y, dtype=float)
    if len(y) > 1:
        dx = np.diff(x)
        area = 0.5 * (y[1:] + y[:-1]) * dx
        out[1:] = np.cumsum(area)
    return out


def reverse_cumulative_trapezoid_np(y, x):
    """Integral from the last coordinate (fixed base) back to each point."""
    y = np.asarray(y, dtype=float)
    x = np.asarray(x, dtype=float)
    out = np.zeros_like(y, dtype=float)
    if len(y) > 1:
        dx = np.diff(x)
        for i in range(len(y) - 2, -1, -1):
            out[i] = out[i + 1] - 0.5 * (y[i + 1] + y[i]) * dx[i]
    return out


# Shear and moment are calculated along the pile/wall length from the free top downward.
# Pressure is psf = lb/ft of wall height per ft wall width.
# Shear output: kips/ft wall width. Moment output: kip-ft/ft wall width.
shear_rankine = cumulative_trapezoid_np(net_pressure_rankine, depths) / 1000.0
shear_coulomb = cumulative_trapezoid_np(net_pressure_coulomb, depths) / 1000.0
moment_rankine = cumulative_trapezoid_np(shear_rankine, depths)
moment_coulomb = cumulative_trapezoid_np(shear_coulomb, depths)


def integrate_curvature_from_fixity(
    curvature,
    coordinate,
    fixed_coordinate,
    restrain_below_fixity=True,
):
    """Integrate curvature upward from a virtual point of fixity.

    The coordinate array increases from the wall top toward the pile toe. At
    the selected point of fixity, rotation and deflection are both zero.

    For the Shoring Suite-style virtual-fixity idealization, the embedded wall
    below the fixity point is treated as restrained; rotation and deflection
    are therefore reported as zero below that point. This avoids extending the
    free-beam curvature integration into the restrained soil zone.

    When ``restrain_below_fixity`` is False, curvature is also integrated from
    the fixity point toward the pile toe. That option is retained for checking,
    but it is not the default virtual-fixity representation.
    """
    curvature = np.asarray(curvature, dtype=float)
    coordinate = np.asarray(coordinate, dtype=float)
    fixed_coordinate = float(np.clip(fixed_coordinate, coordinate[0], coordinate[-1]))

    close = np.isclose(coordinate, fixed_coordinate, rtol=0.0, atol=1.0e-9)
    if np.any(close):
        coordinate_aug = coordinate.copy()
        fixed_index = int(np.flatnonzero(close)[0])
    else:
        coordinate_aug = np.sort(np.append(coordinate, fixed_coordinate))
        fixed_index = int(np.argmin(np.abs(coordinate_aug - fixed_coordinate)))

    curvature_aug = np.interp(coordinate_aug, coordinate, curvature)
    rotation_aug = np.zeros_like(coordinate_aug, dtype=float)
    deflection_aug = np.zeros_like(coordinate_aug, dtype=float)

    # Integrate only the exposed/free portion upward from the virtual fixity.
    for i in range(fixed_index - 1, -1, -1):
        dx = coordinate_aug[i + 1] - coordinate_aug[i]
        rotation_aug[i] = rotation_aug[i + 1] - 0.5 * (
            curvature_aug[i + 1] + curvature_aug[i]
        ) * dx
        deflection_aug[i] = deflection_aug[i + 1] - 0.5 * (
            rotation_aug[i + 1] + rotation_aug[i]
        ) * dx

    if not restrain_below_fixity:
        # Optional diagnostic continuation below the fixity point.
        for i in range(fixed_index, len(coordinate_aug) - 1):
            dx = coordinate_aug[i + 1] - coordinate_aug[i]
            rotation_aug[i + 1] = rotation_aug[i] + 0.5 * (
                curvature_aug[i + 1] + curvature_aug[i]
            ) * dx
            deflection_aug[i + 1] = deflection_aug[i] + 0.5 * (
                rotation_aug[i + 1] + rotation_aug[i]
            ) * dx
    # Otherwise, the initialized zeros are retained below virtual fixity.

    rotation = np.interp(coordinate, coordinate_aug, rotation_aug)
    deflection = np.interp(coordinate, coordinate_aug, deflection_aug)
    return rotation, deflection


def calculate_cantilever_deflection(moment_kip_ft_per_ft):
    """Return curvature, rotation, and deflection using the selected fixity point.

    The pressure calculation produces moment per foot of wall. Multiplying by
    pile_tributary_width_ft gives the bending moment carried by one pile or the
    selected wall strip. E is entered in ksi (kip/in^2), I in in^4, and the
    integration coordinate is converted to inches.
    """
    depth_in = depths * 12.0
    EI_kip_in2 = max(float(youngs_modulus_ksi) * float(moment_of_inertia_in4), 1.0e-12)
    moment_kip_in = np.asarray(moment_kip_ft_per_ft, dtype=float) * float(pile_tributary_width_ft) * 12.0
    curvature_per_in = moment_kip_in / EI_kip_in2

    rotation_rad, deflection_in = integrate_curvature_from_fixity(
        curvature_per_in,
        depth_in,
        float(point_of_fixity_depth) * 12.0,
    )
    return curvature_per_in, rotation_rad, deflection_in


def calculate_elastic_deflection_on_grid(moment_kip_ft_per_ft, depth_grid_ft, fixed_depth_ft):
    """Return curvature, rotation, and deflection for an arbitrary depth grid."""
    depth_grid_ft = np.asarray(depth_grid_ft, dtype=float)
    depth_in = depth_grid_ft * 12.0
    EI_kip_in2 = max(float(youngs_modulus_ksi) * float(moment_of_inertia_in4), 1.0e-12)
    moment_kip_in = (
        np.asarray(moment_kip_ft_per_ft, dtype=float)
        * float(pile_tributary_width_ft)
        * 12.0
    )
    curvature_per_in = moment_kip_in / EI_kip_in2
    rotation_rad, deflection_in = integrate_curvature_from_fixity(
        curvature_per_in,
        depth_in,
        float(fixed_depth_ft) * 12.0,
    )
    return curvature_per_in, rotation_rad, deflection_in


curvature_rankine, rotation_rankine, deflection_rankine = calculate_cantilever_deflection(moment_rankine)
curvature_coulomb, rotation_coulomb, deflection_coulomb = calculate_cantilever_deflection(moment_coulomb)

max_abs_shear_rankine = float(np.max(np.abs(shear_rankine))) if len(shear_rankine) else 0.0
max_abs_shear_coulomb = float(np.max(np.abs(shear_coulomb))) if len(shear_coulomb) else 0.0
max_abs_moment_rankine = float(np.max(np.abs(moment_rankine))) if len(moment_rankine) else 0.0
max_abs_moment_coulomb = float(np.max(np.abs(moment_coulomb))) if len(moment_coulomb) else 0.0
max_abs_deflection_rankine = float(np.max(np.abs(deflection_rankine))) if len(deflection_rankine) else 0.0
max_abs_deflection_coulomb = float(np.max(np.abs(deflection_coulomb))) if len(deflection_coulomb) else 0.0
max_deflection_depth_rankine = float(depths[int(np.argmax(np.abs(deflection_rankine)))]) if len(deflection_rankine) else 0.0
max_deflection_depth_coulomb = float(depths[int(np.argmax(np.abs(deflection_coulomb)))]) if len(deflection_coulomb) else 0.0
top_deflection_rankine = float(deflection_rankine[0]) if len(deflection_rankine) else 0.0
top_deflection_coulomb = float(deflection_coulomb[0]) if len(deflection_coulomb) else 0.0


# -----------------------------
# USACE EM 1110-2-2504 cantilever-wall equilibrium method
# -----------------------------
def usace_pressure_components(z_arr, method):
    """Return USACE net-active and full-net-passive pressure arrays.

    Positive pressure acts toward the excavation/dredge side. In accordance
    with EM 1110-2-2504, passive strengths are reduced by FSp:
    tan(phi_eff) = tan(phi) / FSp and c_eff = c / FSp.
    """
    z_arr = np.asarray(z_arr, dtype=float)
    sigma_eff_retained = np.zeros_like(z_arr)
    water_retained = np.zeros_like(z_arr)
    sigma_eff_dredge = np.zeros_like(z_arr)
    water_dredge = np.zeros_like(z_arr)
    phi_values = np.zeros_like(z_arr)
    cohesion_values = np.zeros_like(z_arr)

    for i, z in enumerate(z_arr):
        _, u_r, sig_r = vertical_stresses_at_depth(float(z))
        _, u_d, sig_d = vertical_stresses_between_depths(
            float(passive_start_depth), float(z)
        )
        phi_i, c_i = properties_at_depth(float(z))
        sigma_eff_retained[i] = sig_r
        water_retained[i] = u_r if show_total_pressure else 0.0
        sigma_eff_dredge[i] = sig_d
        water_dredge[i] = u_d if show_total_pressure else 0.0
        phi_values[i] = phi_i
        cohesion_values[i] = c_i if show_cohesion else 0.0

    passive_phi = np.arctan(
        np.tan(phi_values) / max(float(usace_passive_fs), 1.0e-9)
    )
    passive_c = cohesion_values / max(float(usace_passive_fs), 1.0e-9)

    if method == "Rankine":
        Ka_active = np.array([rankine_Ka(phi, beta) for phi in phi_values])
        Ka_passive_basis = np.array(
            [rankine_Ka(phi_eff, beta) for phi_eff in passive_phi]
        )
    elif method == "Coulomb":
        Ka_active = np.array(
            [coulomb_Ka(phi, delta, alpha, beta) for phi in phi_values]
        )
        Ka_active_fallback = np.array(
            [rankine_Ka(phi, beta) for phi in phi_values]
        )
        Ka_active = np.where(np.isfinite(Ka_active), Ka_active, Ka_active_fallback)

        Ka_passive_basis = np.array(
            [coulomb_Ka(phi_eff, delta, alpha, beta) for phi_eff in passive_phi]
        )
        Ka_passive_fallback = np.array(
            [rankine_Ka(phi_eff, beta) for phi_eff in passive_phi]
        )
        Ka_passive_basis = np.where(
            np.isfinite(Ka_passive_basis),
            Ka_passive_basis,
            Ka_passive_fallback,
        )
    else:
        raise ValueError("USACE method must be 'Rankine' or 'Coulomb'.")

    Ka_active = np.clip(Ka_active, 1.0e-9, None)
    Kp_reduced = 1.0 / np.clip(Ka_passive_basis, 1.0e-9, None)

    retained_active_eff = np.clip(
        Ka_active * sigma_eff_retained
        - 2.0 * cohesion_values * safe_sqrt(Ka_active),
        0.0,
        None,
    )
    dredge_active_eff = np.clip(
        Ka_active * sigma_eff_dredge
        - 2.0 * cohesion_values * safe_sqrt(Ka_active),
        0.0,
        None,
    )
    retained_passive_eff = (
        Kp_reduced * sigma_eff_retained
        + 2.0 * passive_c * safe_sqrt(Kp_reduced)
    )
    dredge_passive_eff = (
        Kp_reduced * sigma_eff_dredge
        + 2.0 * passive_c * safe_sqrt(Kp_reduced)
    )

    dredge_mask = z_arr >= float(passive_start_depth)
    dredge_active_total = np.where(
        dredge_mask,
        dredge_active_eff + water_dredge,
        0.0,
    )
    dredge_passive_total = np.where(
        dredge_mask,
        dredge_passive_eff + water_dredge,
        0.0,
    )
    retained_active_total = retained_active_eff + water_retained
    retained_passive_total = retained_passive_eff + water_retained

    q_for_surcharge = q_uniform if surcharge_type == "Uniform" else q_strip
    usace_surcharge_active = surcharge_pressure(
        z_arr,
        surcharge_type,
        q=q_for_surcharge,
        Q=selected_Q,
        x=selected_x,
        K_arr=Ka_active,
        x1=x_strip_near,
        width=strip_width,
        aashto_load_type=aashto_load_type,
        aashto_Pv=aashto_Pv,
        aashto_bf=aashto_bf,
        aashto_L=aashto_L,
        aashto_d=aashto_d,
    )
    usace_surcharge_passive = surcharge_pressure(
        z_arr,
        surcharge_type,
        q=q_for_surcharge,
        Q=selected_Q,
        x=selected_x,
        K_arr=Kp_reduced,
        x1=x_strip_near,
        width=strip_width,
        aashto_load_type=aashto_load_type,
        aashto_Pv=aashto_Pv,
        aashto_bf=aashto_bf,
        aashto_L=aashto_L,
        aashto_d=aashto_d,
    )

    net_active = (
        retained_active_total
        - dredge_passive_total
        + usace_surcharge_active
    )
    net_passive = (
        retained_passive_total
        - dredge_active_total
        + usace_surcharge_passive
    )

    return {
        "net_active": net_active,
        "net_passive": net_passive,
        "Ka_active": Ka_active,
        "Kp_reduced": Kp_reduced,
    }


def usace_trial_solution(
    embedment_ft,
    transition_fraction,
    method,
    point_count=241,
    master_depths=None,
    master_components=None,
):
    """Evaluate one USACE trial embedment and transition location."""
    embedment_ft = float(embedment_ft)
    transition_fraction = float(transition_fraction)
    bottom_depth = float(passive_start_depth) + embedment_ft
    transition_depth = (
        float(passive_start_depth) + transition_fraction * embedment_ft
    )

    z = np.linspace(0.0, bottom_depth, max(int(point_count), 81))

    # The earth-pressure components depend on depth and the selected method,
    # but not on the trial embedment or transition ratio. During the USACE
    # search they are therefore calculated once on a master depth grid and
    # interpolated for each trial. This removes thousands of repeated layered-
    # stress calculations and keeps the Streamlit app responsive.
    if master_depths is not None and master_components is not None:
        master_depths = np.asarray(master_depths, dtype=float)
        net_active = np.interp(
            z, master_depths, np.asarray(master_components["net_active"], dtype=float)
        )
        net_passive = np.interp(
            z, master_depths, np.asarray(master_components["net_passive"], dtype=float)
        )
    else:
        components = usace_pressure_components(z, method)
        net_active = components["net_active"]
        net_passive = components["net_passive"]

    p_transition = float(np.interp(transition_depth, z, net_active))
    p_bottom_passive = float(net_passive[-1])
    denominator = max(bottom_depth - transition_depth, 1.0e-9)
    lower_linear = p_transition + (
        p_bottom_passive - p_transition
    ) * (z - transition_depth) / denominator
    design_pressure = np.where(z <= transition_depth, net_active, lower_linear)

    force_residual = float(np.trapezoid(design_pressure, z))
    moment_residual_about_bottom = float(
        np.trapezoid(design_pressure * (bottom_depth - z), z)
    )
    pressure_scale = max(float(np.max(np.abs(design_pressure))), 1.0)
    normalized_residual = np.array(
        [
            force_residual / (pressure_scale * max(bottom_depth, 1.0)),
            moment_residual_about_bottom
            / (pressure_scale * max(bottom_depth, 1.0) ** 2),
        ],
        dtype=float,
    )

    return {
        "depths": z,
        "net_active": net_active,
        "net_passive": net_passive,
        "design_pressure": design_pressure,
        "bottom_depth": bottom_depth,
        "embedment": embedment_ft,
        "transition_depth": transition_depth,
        "transition_fraction": transition_fraction,
        "z_above_bottom": bottom_depth - transition_depth,
        "pressure_at_transition": p_transition,
        "full_net_passive_at_bottom": p_bottom_passive,
        "force_residual": force_residual,
        "moment_residual": moment_residual_about_bottom,
        "normalized_residual": normalized_residual,
    }


def bounded_grid_search_numpy(
    objective,
    x0,
    lower_bounds,
    upper_bounds,
    coarse_points=11,
    refinement_points=7,
    refinement_levels=8,
):
    """Fast bounded two-variable search implemented with NumPy only.

    The USACE equilibrium problem contains only two unknowns: embedment and
    transition-depth ratio. A small global grid is followed by shrinking local
    grids. Objective values are cached, so repeated candidate points do not
    trigger duplicate pressure integrations.
    """
    lower = np.asarray(lower_bounds, dtype=float)
    upper = np.asarray(upper_bounds, dtype=float)
    x0 = np.clip(np.asarray(x0, dtype=float), lower, upper)
    cache = {}
    nfev = 0

    def evaluate(values):
        nonlocal nfev
        values = np.clip(np.asarray(values, dtype=float), lower, upper)
        key = tuple(np.round(values, 12))
        if key in cache:
            return cache[key]
        residual = np.asarray(objective(values), dtype=float)
        nfev += 1
        if residual.shape != (2,) or not np.all(np.isfinite(residual)):
            residual = np.array([1.0e6, 1.0e6], dtype=float)
        cost = 0.5 * float(np.dot(residual, residual))
        result = (cost, values.copy(), residual.copy())
        cache[key] = result
        return result

    candidates = [evaluate(x0)]
    x_values = np.linspace(lower[0], upper[0], max(int(coarse_points), 5))
    y_values = np.linspace(lower[1], upper[1], max(int(coarse_points), 5))
    for x_value in x_values:
        for y_value in y_values:
            candidates.append(evaluate([x_value, y_value]))

    candidates.sort(key=lambda item: item[0])
    # Refine the best global point and two alternate low-cost basins.
    centers = [item[1].copy() for item in candidates[:3]]
    global_best = candidates[0]
    coarse_step = np.array(
        [
            (upper[0] - lower[0]) / max(len(x_values) - 1, 1),
            (upper[1] - lower[1]) / max(len(y_values) - 1, 1),
        ],
        dtype=float,
    )

    for initial_center in centers:
        center = initial_center.copy()
        half_width = np.maximum(1.5 * coarse_step, np.array([0.05, 0.01]))
        local_best = evaluate(center)

        for _ in range(max(int(refinement_levels), 1)):
            x_lo = max(lower[0], center[0] - half_width[0])
            x_hi = min(upper[0], center[0] + half_width[0])
            y_lo = max(lower[1], center[1] - half_width[1])
            y_hi = min(upper[1], center[1] + half_width[1])
            local_candidates = []
            for x_value in np.linspace(x_lo, x_hi, max(int(refinement_points), 5)):
                for y_value in np.linspace(y_lo, y_hi, max(int(refinement_points), 5)):
                    local_candidates.append(evaluate([x_value, y_value]))
            local_candidates.sort(key=lambda item: item[0])
            if local_candidates[0][0] < local_best[0]:
                local_best = local_candidates[0]
                center = local_best[1].copy()
            half_width *= 0.35

        if local_best[0] < global_best[0]:
            global_best = local_best

    return {
        "x": global_best[1],
        "fun": global_best[2],
        "cost": global_best[0],
        "nfev": nfev,
        "success": bool(np.all(np.isfinite(global_best[2]))),
    }

def solve_usace_cantilever(method):
    """Solve USACE force and moment equilibrium for penetration and transition."""
    if not run_usace_analysis:
        return {"success": False, "message": "USACE analysis is disabled."}
    if float(passive_start_depth) <= 0.0:
        return {
            "success": False,
            "message": (
                "The USACE cantilever method requires an excavation/dredge-line depth "
                "greater than 0 ft."
            ),
        }

    min_embedment = max(0.25, 0.025 * float(passive_start_depth))
    max_embedment = max(float(usace_max_embedment_ft), min_embedment + 0.5)

    # Precompute the depth-dependent pressure components once. Trial solutions
    # then use interpolation instead of recalculating layered vertical stresses
    # at every optimization evaluation.
    maximum_bottom_depth = float(passive_start_depth) + max_embedment
    master_point_count = int(
        min(
            3001,
            max(801, 2 * int(usace_solver_points) + 1, 4 * int(n_points) + 1),
        )
    )
    master_depths = np.linspace(0.0, maximum_bottom_depth, master_point_count)
    master_components = usace_pressure_components(master_depths, method)

    def objective(x):
        trial = usace_trial_solution(
            x[0],
            x[1],
            method,
            point_count=121,
            master_depths=master_depths,
            master_components=master_components,
        )
        return trial["normalized_residual"]

    current_embedment = max(float(H) - float(passive_start_depth), min_embedment)
    result = bounded_grid_search_numpy(
        objective,
        x0=[min(max(current_embedment, min_embedment), max_embedment), 0.60],
        lower_bounds=[min_embedment, 0.01],
        upper_bounds=[max_embedment, 0.99],
        coarse_points=11,
        refinement_points=7,
        refinement_levels=8,
    )

    final = usace_trial_solution(
        result["x"][0],
        result["x"][1],
        method,
        point_count=max(int(usace_solver_points), int(n_points), 301),
        master_depths=master_depths,
        master_components=master_components,
    )
    final_residual_norm = float(np.linalg.norm(final["normalized_residual"]))
    at_upper_bound = final["embedment"] >= 0.999 * max_embedment
    success = final_residual_norm <= 5.0e-4 and not at_upper_bound

    if not success:
        message = (
            f"No converged {method} equilibrium solution was found within the maximum "
            f"trial embedment of {max_embedment:.2f} ft. Increase the search limit or "
            "review the pressure inputs."
        )
        final.update(
            {
                "success": False,
                "message": message,
                "residual_norm": final_residual_norm,
                "solver_evaluations": int(result.get("nfev", 0)),
            }
        )
        return final

    shear = cumulative_trapezoid_np(final["design_pressure"], final["depths"]) / 1000.0
    moment = cumulative_trapezoid_np(shear, final["depths"])
    curvature, rotation, deflection = calculate_elastic_deflection_on_grid(
        moment,
        final["depths"],
        final["transition_depth"],
    )

    final.update(
        {
            "success": True,
            "message": "Converged",
            "shear": shear,
            "moment": moment,
            "curvature": curvature,
            "rotation": rotation,
            "deflection": deflection,
            "max_abs_shear": float(np.max(np.abs(shear))),
            "max_abs_moment": float(np.max(np.abs(moment))),
            "top_deflection": float(deflection[0]),
            "max_abs_deflection": float(np.max(np.abs(deflection))),
            "max_deflection_depth": float(
                final["depths"][int(np.argmax(np.abs(deflection)))]
            ),
            "residual_norm": final_residual_norm,
            "solver_evaluations": int(result.get("nfev", 0)),
        }
    )
    return final


if run_usace_analysis:
    with st.spinner("Solving USACE force and moment equilibrium..."):
        usace_rankine = solve_usace_cantilever("Rankine")
        usace_coulomb = solve_usace_cantilever("Coulomb")
else:
    usace_rankine = {"success": False, "message": "USACE analysis is disabled."}
    usace_coulomb = {"success": False, "message": "USACE analysis is disabled."}


# -----------------------------
# Input geometry schematic
# -----------------------------
def create_geometry_figure():
    """Creates a schematic based on the current input geometry, soil layers, water table, and surcharge."""
    fig, ax = plt.subplots(figsize=(11, 6.5))
    fig.patch.set_facecolor("#f8fafc")
    ax.set_facecolor("#f8fafc")

    h = float(H)
    tan_alpha = np.tan(alpha)
    tan_beta = np.tan(beta)
    wall_x_top = 0.0
    wall_x_bottom = h * tan_alpha

    max_surcharge_x = 0.0
    if surcharge_type == "Line Load":
        max_surcharge_x = float(x_line)
    elif surcharge_type == "Point Load":
        max_surcharge_x = float(x_point)
    elif surcharge_type == "Strip Load":
        max_surcharge_x = float(x_strip_near) + float(strip_width)
    elif surcharge_type == "AASHTO":
        max_surcharge_x = float(aashto_d) + max(float(aashto_bf), 0.0) / 2.0

    x_right = max(h * 1.25, max_surcharge_x + h * 0.25, 20.0)
    y_surface_right = -x_right * tan_beta
    y_top = min(-2.0, y_surface_right - 2.0)
    y_bottom = h + max(2.0, 0.08 * h)

    # Soil mass behind the wall.
    wall_x = np.array([wall_x_top, wall_x_bottom])
    wall_y = np.array([0.0, h])
    soil_poly_x = [wall_x_top, x_right, x_right, wall_x_bottom]
    soil_poly_y = [0.0, y_surface_right, h, h]
    ax.fill(soil_poly_x, soil_poly_y, color="#d97706", alpha=0.12, label="Backfill soil")

    # Front side excavation / passive side.
    # The ground in front of the wall is shown only below the passive/excavation start depth.
    # Above this elevation the front side is left blank to represent the excavation cut.
    x_left = -0.12 * h
    exc_depth = max(0.0, min(float(passive_start_depth), h))
    if exc_depth < h:
        ax.fill(
            [x_left, exc_depth * tan_alpha, wall_x_bottom, x_left],
            [exc_depth, exc_depth, h, h],
            color="#e5e7eb",
            alpha=0.45,
            ec="none",
            label="Excavation / passive side",
        )
        ax.plot([x_left, exc_depth * tan_alpha], [exc_depth, exc_depth], color="#111827", lw=2.4)

    # Wall line and thickness.
    ax.plot(wall_x, wall_y, color="#334155", lw=5, solid_capstyle="round", label="Wall")
    ax.plot(wall_x - 0.25, wall_y, color="#94a3b8", lw=5, alpha=0.9)

    # Backfill surface.
    ax.plot([0, x_right], [0, y_surface_right], color="#78350f", lw=2.5, label=f"Backfill slope β = {beta_deg:.2f}°")

    # Layer boundaries and labels. Use different fill colors so each soil layer is visible.
    top_depth = 0.0
    layer_colors = [
        "#fef3c7", "#dbeafe", "#dcfce7", "#fee2e2",
        "#ede9fe", "#ffedd5", "#cffafe", "#fce7f3"
    ]

    def draw_front_excavation_layer(i, y1, y2):
        """Show passive-side soil below the excavation/passive start depth."""
        y1 = max(y1, exc_depth)
        y2 = min(y2, h)
        if y2 <= y1:
            return
        ax.fill(
            [x_left, y1 * tan_alpha, y2 * tan_alpha, x_left],
            [y1, y1, y2, y2],
            color=layer_colors[i % len(layer_colors)],
            alpha=0.30,
            ec="none",
        )

    def draw_soil_layer(i, row, y1, y2, extended=False):
        """Draw one layer interval and label its soil parameters."""
        x_wall_y1 = y1 * tan_alpha
        x_wall_y2 = y2 * tan_alpha
        ax.fill(
            [x_wall_y1, x_right, x_right, x_wall_y2],
            [y1, y1, y2, y2],
            color=layer_colors[i % len(layer_colors)],
            alpha=0.42,
            ec="#e2e8f0",
            lw=0.6,
            label="Soil layer" if i == 0 else None,
        )
        if y2 < h:
            ax.hlines(y2, xmin=y2 * tan_alpha, xmax=x_right, colors="#64748b", linestyles="--", lw=1.0)

        mid = 0.5 * (y1 + y2)
        gamma_m = float(row["Moist Unit Weight γm (pcf)"])
        gamma_sat = float(row["Saturated Unit Weight γsat (pcf)"])
        phi_i = float(row["Friction Angle φ (deg)"])
        c_i = float(row["Cohesion c (psf)"])
        label_suffix = " (extended)" if extended else ""
        ax.text(
            x_right * 0.72,
            mid,
            f"Layer {i + 1}{label_suffix}\nγm={gamma_m:.0f} pcf, γsat={gamma_sat:.0f} pcf\nφ={phi_i:.0f}°, c={c_i:.0f} psf",
            fontsize=8,
            color="#334155",
            ha="left",
            va="center",
            bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#cbd5e1", alpha=0.82),
        )

    last_row = None
    last_i = 0
    for i, row in layer_df.iterrows():
        bottom_depth = min(float(row["Bottom Depth (ft)"]), h)
        if bottom_depth <= top_depth:
            continue
        draw_soil_layer(i, row, top_depth, bottom_depth)
        draw_front_excavation_layer(i, top_depth, bottom_depth)
        top_depth = bottom_depth
        last_row = row
        last_i = i

    # If the final entered layer bottom is above the wall base, show that
    # the final layer's properties are extended to the base for calculation.
    if top_depth < h and last_row is not None:
        draw_soil_layer(last_i, last_row, top_depth, h, extended=True)
        draw_front_excavation_layer(last_i, top_depth, h)

    # Water table.
    if include_water:
        ax.hlines(float(water_table), xmin=float(water_table) * tan_alpha, xmax=x_right,
                  colors="#0284c7", linestyles=":", lw=2.0, label=f"Water table = {water_table:.2f} ft")
        ax.text(x_right * 0.04 + float(water_table) * tan_alpha, float(water_table) - 0.35,
                f"WT {water_table:.2f} ft", color="#0284c7", fontsize=9, fontweight="bold")
        ax.fill_between([float(water_table) * tan_alpha, x_right], float(water_table), h,
                        color="#38bdf8", alpha=0.10)

    # Surcharge geometry.
    if surcharge_type == "Uniform":
        xs = np.linspace(1.0, min(x_right - 1.0, h), 5)
        for x in xs:
            y = -x * tan_beta
            ax.annotate("", xy=(x, y + 0.1), xytext=(x, y - 1.6),
                        arrowprops=dict(arrowstyle="-|>", color="#dc2626", lw=1.8))
        ax.text(min(x_right * 0.18, h), y_top + 0.5, f"Uniform surcharge q = {q_uniform:.0f} psf",
                color="#dc2626", fontsize=9, fontweight="bold")
    elif surcharge_type == "Line Load":
        x = float(x_line)
        y = -x * tan_beta
        ax.annotate("", xy=(x, y + 0.1), xytext=(x, y - 2.6),
                    arrowprops=dict(arrowstyle="-|>", color="#dc2626", lw=2.2))
        ax.text(x, y - 2.9, f"Line load\nQ={Q_line:.0f} lb/ft\nx={x_line:.2f} ft", color="#dc2626",
                fontsize=9, ha="center", va="top", bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#fecaca", alpha=0.9))
    elif surcharge_type == "Point Load":
        x = float(x_point)
        y = -x * tan_beta
        ax.annotate("", xy=(x, y + 0.1), xytext=(x, y - 2.6),
                    arrowprops=dict(arrowstyle="-|>", color="#dc2626", lw=2.2))
        ax.text(x, y - 2.9, f"Point load\nQ={Q_point:.0f} lb\nx={x_point:.2f} ft", color="#dc2626",
                fontsize=9, ha="center", va="top", bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#fecaca", alpha=0.9))
    elif surcharge_type == "Strip Load":
        x1 = float(x_strip_near)
        x2 = float(x_strip_near + strip_width)
        y1 = -x1 * tan_beta
        y2 = -x2 * tan_beta
        ax.plot([x1, x2], [y1, y2], color="#dc2626", lw=5, solid_capstyle="butt")
        for x in np.linspace(x1, x2, 5):
            y = -x * tan_beta
            ax.annotate("", xy=(x, y + 0.1), xytext=(x, y - 1.5),
                        arrowprops=dict(arrowstyle="-|>", color="#dc2626", lw=1.6))
        ax.text((x1 + x2) / 2.0, min(y1, y2) - 2.0,
                f"Strip load q={q_strip:.0f} psf\nx1={x1:.2f} ft, B={strip_width:.2f} ft",
                color="#dc2626", fontsize=9, ha="center", va="top",
                bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#fecaca", alpha=0.9))
    elif surcharge_type == "AASHTO":
        x = float(aashto_d)
        y = -x * tan_beta
        width_plot = max(float(aashto_bf), 0.5)
        if aashto_load_type == "Point load":
            ax.plot(x, y, marker="o", color="#dc2626", ms=8)
        else:
            ax.plot([x - width_plot / 2.0, x + width_plot / 2.0], [y, y], color="#dc2626", lw=5, solid_capstyle="butt")
        ax.annotate("", xy=(x, y + 0.1), xytext=(x, y - 2.6),
                    arrowprops=dict(arrowstyle="-|>", color="#dc2626", lw=2.2))
        ax.text(x, y - 2.9,
                f"AASHTO {aashto_load_type}\nP={aashto_Pv:.0f} lb{'/ft' if aashto_load_type == 'Strip footing' else ''}\nd={aashto_d:.2f} ft" + (f", bf={aashto_bf:.2f} ft" if aashto_load_type != "Point load" else ", bf=0"),
                color="#dc2626", fontsize=9, ha="center", va="top",
                bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#fecaca", alpha=0.9))

    # Dimension arrows for wall height.
    ax.annotate("", xy=(-1.2, 0), xytext=(-1.2, h),
                arrowprops=dict(arrowstyle="<->", color="#0f172a", lw=1.2))
    ax.text(-1.6, h / 2.0, f"H = {h:.2f} ft", rotation=90, ha="center", va="center", fontsize=9, color="#0f172a")

    # No title inside the figure; the app/report section heading provides context.
    ax.set_xlabel("Horizontal distance from wall top (ft)")
    ax.set_ylabel("Depth from top (ft)")
    ax.set_xlim(-0.12 * h, x_right)
    ax.set_ylim(y_bottom, y_top)
    ax.grid(True, alpha=0.25, linestyle="--")
    ax.legend(loc="lower right", fontsize=8)
    fig.tight_layout()
    return fig


# -----------------------------
# Report generation helpers
# -----------------------------
def fig_to_png_bytes(fig, dpi=180):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    return buf


def create_pressure_diagram_figure():
    fig, axes = plt.subplots(1, 3, figsize=(12, 7), sharey=True)
    configs = [
        ("Rankine", pa_rankine, pp_rankine),
        ("Coulomb", pa_coulomb, pp_coulomb),
        ("At-Rest", pa_atrest, pa_atrest),
    ]
    for ax, (title, pa, pp) in zip(axes, configs):
        ax.plot(-pa, depths, lw=2.0, label="Active / At-Rest")
        ax.plot(pp, depths, lw=2.0, ls="--", label="Passive")
        ax.axvline(0, lw=0.8)
        ax.set_title(title)
        ax.set_xlabel("Pressure (psf)")
        ax.grid(True, alpha=0.3, linestyle="--")
        ax.legend(fontsize=8)
    # Shared depth axis: zero at wall top and depth increasing downward.
    axes[0].set_ylim(H, 0)
    axes[0].set_ylabel("Depth (ft)")
    fig.suptitle("Lateral Earth Pressure Diagrams")
    fig.tight_layout()
    return fig


def create_active_comparison_figure():
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.plot(pa_rankine, depths, lw=2.2, label="Rankine Active")
    ax.plot(pa_coulomb, depths, lw=2.2, ls="--", label="Coulomb Active")
    ax.plot(pa_atrest, depths, lw=2.2, ls=":", label="At-Rest")
    if surcharge_type != "None":
        ax.plot(pa_rankine_plus_surcharge, depths, lw=1.8, ls="-.", label="Rankine + surcharge")
    ax.set_ylim(H, 0)
    ax.set_xlabel("Pressure (psf)")
    ax.set_ylabel("Depth (ft)")
    ax.set_title("Active / At-Rest Pressure Comparison")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig


def create_passive_comparison_figure():
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.plot(pp_rankine, depths, lw=2.2, label="Rankine Passive")
    ax.plot(pp_coulomb, depths, lw=2.2, ls="--", label="Coulomb Passive")
    ax.set_ylim(H, 0)
    ax.set_xlabel("Pressure (psf)")
    ax.set_ylabel("Depth (ft)")
    ax.set_title("Passive Pressure Comparison")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig


def create_surcharge_figure():
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.plot(sur_r, depths, lw=2.2, label="Rankine/rigid")
    if surcharge_type in ["Uniform", "AASHTO"]:
        ax.plot(sur_c, depths, lw=2.2, ls="--", label="Coulomb")
        ax.plot(sur_0, depths, lw=2.2, ls=":", label="At-Rest")
    ax.set_ylim(H, 0)
    ax.set_xlabel("Surcharge Pressure (psf)")
    ax.set_ylabel("Depth (ft)")
    ax.set_title("Separate Surcharge Pressure")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig


def create_net_shear_moment_figure():
    fig, axes = plt.subplots(1, 4, figsize=(19, 6), sharey=True)

    axes[0].plot(net_pressure_rankine, depths, lw=2.2, label="Rankine net")
    axes[0].plot(net_pressure_coulomb, depths, lw=2.2, ls="--", label="Coulomb net")
    axes[0].axvline(0, color="#334155", lw=0.8)
    axes[0].set_xlabel("Net pressure (psf)")
    axes[0].set_ylabel("Depth from top (ft)")
    axes[0].set_title("Net Wall Pressure")

    axes[1].plot(shear_rankine, depths, lw=2.2, label="Rankine")
    axes[1].plot(shear_coulomb, depths, lw=2.2, ls="--", label="Coulomb")
    axes[1].axvline(0, color="#334155", lw=0.8)
    axes[1].set_xlabel("Shear V (kips/ft)")
    axes[1].set_title("Shear Diagram")

    axes[2].plot(moment_rankine, depths, lw=2.2, label="Rankine")
    axes[2].plot(moment_coulomb, depths, lw=2.2, ls="--", label="Coulomb")
    axes[2].axvline(0, color="#334155", lw=0.8)
    axes[2].set_xlabel("Moment M (kip-ft/ft)")
    axes[2].set_title("Moment Diagram")

    axes[3].plot(deflection_rankine, depths, lw=2.2, label="Rankine")
    axes[3].plot(deflection_coulomb, depths, lw=2.2, ls="--", label="Coulomb")
    axes[3].axvline(0, color="#334155", lw=0.8)
    axes[3].axhline(
        point_of_fixity_depth,
        color="#7c3aed",
        lw=1.1,
        ls="-.",
        label=f"Fixity = {point_of_fixity_depth:.2f} ft",
    )
    axes[3].set_xlabel("Deflection y (in)")
    axes[3].set_title("Elastic Deflection")

    for ax in axes:
        ax.grid(True, alpha=0.3, linestyle="--")
        ax.legend(fontsize=8, loc="upper right")

        # Place the horizontal-axis ticks and labels at the top of each panel.
        ax.xaxis.set_label_position("top")
        ax.xaxis.tick_top()
        ax.tick_params(axis="x", labeltop=True, labelbottom=False, top=True, bottom=False)
        ax.xaxis.labelpad = 8

        # Keep the panel title above the top axis label and tick values.
        ax.title.set_position((0.5, 1.20))

    # Because the four panels share the y-axis, set the direction once only.
    # This guarantees 0 ft at the top and H at the bottom.
    axes[0].set_ylim(H, 0)
    fig.tight_layout(rect=[0.0, 0.0, 1.0, 0.93])
    return fig


def create_usace_cantilever_figure():
    """Create USACE pressure, shear, moment, and approximate deflection diagrams."""
    fig, axes = plt.subplots(2, 4, figsize=(19, 11), sharey="row")
    solutions = [("Rankine", usace_rankine), ("Coulomb", usace_coulomb)]

    for row, (method, solution) in enumerate(solutions):
        row_axes = axes[row]
        if not solution.get("success", False):
            for ax in row_axes:
                ax.axis("off")
                ax.text(
                    0.5,
                    0.5,
                    solution.get("message", "USACE solution unavailable."),
                    ha="center",
                    va="center",
                    wrap=True,
                    transform=ax.transAxes,
                )
            continue

        z = solution["depths"]
        bottom = solution["bottom_depth"]
        transition = solution["transition_depth"]

        row_axes[0].plot(
            solution["design_pressure"], z, lw=2.3, label="USACE design"
        )
        row_axes[0].plot(
            solution["net_active"], z, lw=1.5, ls="--", label="Net active"
        )
        row_axes[0].plot(
            solution["net_passive"], z, lw=1.5, ls=":", label="Full net passive"
        )
        row_axes[0].axvline(0.0, color="#334155", lw=0.8)
        row_axes[0].set_xlabel("Pressure (psf)")
        row_axes[0].set_ylabel("Depth from top (ft)")
        row_axes[0].set_title(f"{method}: USACE Pressure")

        row_axes[1].plot(solution["shear"], z, lw=2.2, label=method)
        row_axes[1].axvline(0.0, color="#334155", lw=0.8)
        row_axes[1].set_xlabel("Shear V (kips/ft)")
        row_axes[1].set_title(
            f"Shear |V|max={solution['max_abs_shear']:.2f}"
        )

        row_axes[2].plot(solution["moment"], z, lw=2.2, label=method)
        row_axes[2].axvline(0.0, color="#334155", lw=0.8)
        row_axes[2].set_xlabel("Moment M (kip-ft/ft)")
        row_axes[2].set_title(
            f"Moment |M|max={solution['max_abs_moment']:.2f}"
        )

        row_axes[3].plot(solution["deflection"], z, lw=2.2, label=method)
        row_axes[3].axvline(0.0, color="#334155", lw=0.8)
        row_axes[3].set_xlabel("Deflection y (in)")
        row_axes[3].set_title(
            f"Approx. Deflection |y|max={solution['max_abs_deflection']:.3f}"
        )

        for ax in row_axes:
            ax.axhline(
                float(passive_start_depth),
                color="#92400e",
                lw=1.0,
                ls="--",
                label="Excavation / dredge line" if ax is row_axes[0] else None,
            )
            ax.axhline(
                transition,
                color="#7c3aed",
                lw=1.1,
                ls="-.",
                label="USACE transition" if ax is row_axes[0] else None,
            )
            ax.set_ylim(bottom, 0.0)
            ax.grid(True, alpha=0.3, linestyle="--")
            ax.legend(fontsize=8, loc="upper right")
            ax.xaxis.set_label_position("top")
            ax.xaxis.tick_top()
            ax.tick_params(
                axis="x",
                labeltop=True,
                labelbottom=False,
                top=True,
                bottom=False,
            )
            ax.xaxis.labelpad = 8
            ax.title.set_position((0.5, 1.20))

    fig.suptitle(
        "USACE EM 1110-2-2504 Cantilever-Wall Equilibrium Results",
        fontsize=14,
        fontweight="bold",
        y=0.995,
    )
    fig.tight_layout(rect=[0.0, 0.0, 1.0, 0.96])
    return fig


def usace_summary_df():
    rows = []
    for method, solution in [("Rankine", usace_rankine), ("Coulomb", usace_coulomb)]:
        if not solution.get("success", False):
            rows.append(
                {
                    "Method": method,
                    "Status": solution.get("message", "No solution"),
                    "Required penetration d (ft)": "--",
                    "Design bottom depth (ft)": "--",
                    "Transition depth from top (ft)": "--",
                    "z above bottom (ft)": "--",
                    "Max |V| (kips/ft)": "--",
                    "Max |M| (kip-ft/ft)": "--",
                    "Top deflection (in)": "--",
                    "Max |deflection| (in)": "--",
                    "Force residual (lb/ft)": "--",
                    "Moment residual (lb-ft/ft)": "--",
                }
            )
            continue
        rows.append(
            {
                "Method": method,
                "Status": "Converged",
                "Required penetration d (ft)": f"{solution['embedment']:.2f}",
                "Design bottom depth (ft)": f"{solution['bottom_depth']:.2f}",
                "Transition depth from top (ft)": f"{solution['transition_depth']:.2f}",
                "z above bottom (ft)": f"{solution['z_above_bottom']:.2f}",
                "Max |V| (kips/ft)": f"{solution['max_abs_shear']:.2f}",
                "Max |M| (kip-ft/ft)": f"{solution['max_abs_moment']:.2f}",
                "Top deflection (in)": f"{solution['top_deflection']:.3f}",
                "Max |deflection| (in)": f"{solution['max_abs_deflection']:.3f}",
                "Force residual (lb/ft)": f"{solution['force_residual']:.3f}",
                "Moment residual (lb-ft/ft)": f"{solution['moment_residual']:.3f}",
            }
        )
    return pd.DataFrame(rows)


def usace_detail_df(solution):
    if not solution.get("success", False):
        return pd.DataFrame()
    idx = np.linspace(
        0,
        len(solution["depths"]) - 1,
        min(60, len(solution["depths"])),
        dtype=int,
    )
    return pd.DataFrame(
        {
            "Depth (ft)": np.round(solution["depths"][idx], 3),
            "Net active pressure (psf)": np.round(solution["net_active"][idx], 2),
            "Full net passive pressure (psf)": np.round(solution["net_passive"][idx], 2),
            "USACE design pressure (psf)": np.round(solution["design_pressure"][idx], 2),
            "Shear (kips/ft)": np.round(solution["shear"][idx], 4),
            "Moment (kip-ft/ft)": np.round(solution["moment"][idx], 4),
            "Rotation (rad)": np.round(solution["rotation"][idx], 7),
            "Deflection (in)": np.round(solution["deflection"][idx], 5),
        }
    )


def report_table_df():
    idx = np.linspace(0, len(depths) - 1, min(30, len(depths)), dtype=int)
    return pd.DataFrame({
        "Depth (ft)": np.round(depths[idx], 2),
        "Layer": [layer_index_at_depth(float(z)) + 1 for z in depths[idx]],
        "sigma_v eff (psf)": np.round(sigma_v_eff[idx], 1),
        "u water (psf)": np.round(u_water[idx], 1),
        "Ka Rankine": np.round(Ka_r[idx], 3),
        "Ka Coulomb": np.round(Ka_c[idx], 3),
        "Rankine active (psf)": np.round(pa_rankine[idx], 1),
        "Coulomb active (psf)": np.round(pa_coulomb[idx], 1),
        "At-rest (psf)": np.round(pa_atrest[idx], 1),
        "Surcharge R (psf)": np.round(sur_r[idx], 1),
        "Rankine passive effective (psf)": np.round(pp_rankine_eff[idx], 1),
        "Coulomb passive effective (psf)": np.round(pp_coulomb_eff[idx], 1),
        "Rankine passive total (psf)": np.round(pp_rankine[idx], 1),
        "Coulomb passive total (psf)": np.round(pp_coulomb[idx], 1),
        "Rankine net pressure (psf)": np.round(net_pressure_rankine[idx], 1),
        "Coulomb net pressure (psf)": np.round(net_pressure_coulomb[idx], 1),
        "Rankine shear (kips/ft)": np.round(shear_rankine[idx], 3),
        "Coulomb shear (kips/ft)": np.round(shear_coulomb[idx], 3),
        "Rankine moment (kip-ft/ft)": np.round(moment_rankine[idx], 3),
        "Coulomb moment (kip-ft/ft)": np.round(moment_coulomb[idx], 3),
        "Rankine rotation (rad)": np.round(rotation_rankine[idx], 6),
        "Coulomb rotation (rad)": np.round(rotation_coulomb[idx], 6),
        "Rankine deflection (in)": np.round(deflection_rankine[idx], 4),
        "Coulomb deflection (in)": np.round(deflection_coulomb[idx], 4),
    })


def summary_table_df():
    return pd.DataFrame({
        "Method": ["Rankine", "Coulomb", "At-Rest"],
        "Earth/Water resultant (kips/ft)": [f"{Fa_r:.2f}", f"{Fa_c:.2f}", f"{Fa_0:.2f}"],
        "Height above base (ft)": [f"{ha_r:.2f}", f"{ha_c:.2f}", f"{ha_0:.2f}"],
        "Surcharge resultant (kips/ft)": [f"{Fs_r:.2f}", f"{Fs_c:.2f}", f"{Fs_0:.2f}"],
        "Combined resultant (kips/ft)": [f"{Ft_r:.2f}", f"{Ft_c:.2f}", f"{Ft_0:.2f}"],
        "Passive resultant (kips/ft)": [f"{Fp_r:.2f}", f"{Fp_c:.2f}", "--"],
        "Max abs shear (kips/ft)": [f"{max_abs_shear_rankine:.2f}", f"{max_abs_shear_coulomb:.2f}", "--"],
        "Max abs moment (kip-ft/ft)": [f"{max_abs_moment_rankine:.2f}", f"{max_abs_moment_coulomb:.2f}", "--"],
        "Top deflection (in)": [f"{top_deflection_rankine:.3f}", f"{top_deflection_coulomb:.3f}", "--"],
        "Max abs deflection (in)": [f"{max_abs_deflection_rankine:.3f}", f"{max_abs_deflection_coulomb:.3f}", "--"],
        "Point of fixity depth (ft)": [f"{point_of_fixity_depth:.2f}", f"{point_of_fixity_depth:.2f}", "--"],
    })


def add_df_to_doc(doc, df, max_rows=None):
    if max_rows is not None:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
    return table


def generate_word_report(template_bytes=None, meta=None):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    meta = meta or {}
    template_name = "TSG Technical MEMO 2026.docx"
    if template_bytes is not None:
        doc = Document(BytesIO(template_bytes))
    elif os.path.exists(template_name):
        doc = Document(template_name)
    elif os.path.exists(os.path.join(os.getcwd(), template_name)):
        doc = Document(os.path.join(os.getcwd(), template_name))
    elif os.path.exists(os.path.join("/mnt/data", template_name)):
        doc = Document(os.path.join("/mnt/data", template_name))
    else:
        doc = Document()

    doc.add_page_break()
    title = doc.add_heading("Lateral Earth Pressure and Deflection Calculation Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Date: {meta.get('date', date.today().isoformat())}").bold = True

    doc.add_heading("Memorandum Information", level=2)
    info = pd.DataFrame({
        "Item": ["To", "Cc", "From", "Project", "Subject"],
        "Value": [meta.get("to", ""), meta.get("cc", ""), meta.get("from", ""), meta.get("project", ""), meta.get("subject", "")],
    })
    add_df_to_doc(doc, info)

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This report summarizes the lateral earth pressure calculation for the wall geometry, layered backfill, groundwater condition, passive pressure zone, and surcharge loading entered in the Streamlit calculator. "
        "The calculation reports active, at-rest, passive, surcharge-only, and combined lateral pressure diagrams in English units. "
        "It also estimates elastic pile/wall rotation and deflection from the calculated bending-moment diagram using the entered E and I values."
    )

    doc.add_heading("Input Parameters", level=2)
    input_df = pd.DataFrame({
        "Parameter": ["Wall height H", "Wall inclination alpha", "Backfill slope beta", "Wall friction delta", "Passive pressure start depth", "Water table", "Water unit weight", "Surcharge type", "Young's modulus E", "Moment of inertia I", "Pile tributary width / spacing", "Deflection boundary condition", "Embedment depth", "Point of fixity depth", "Run USACE analysis", "USACE passive FSp", "USACE maximum trial embedment"],
        "Value": [f"{H:.2f} ft", f"{alpha_deg:.2f} deg", f"{beta_deg:.2f} deg", f"{delta_deg:.2f} deg", f"{passive_start_depth:.2f} ft", f"{water_table:.2f} ft" if include_water else "Not included", f"{gamma_w:.1f} pcf" if include_water else "N/A", surcharge_label, f"{youngs_modulus_ksi:.1f} ksi", f"{moment_of_inertia_in4:.2f} in^4", f"{pile_tributary_width_ft:.2f} ft", deflection_boundary_label, f"{embedment_depth_ft:.2f} ft", f"{point_of_fixity_depth:.2f} ft", "Yes" if run_usace_analysis else "No", f"{usace_passive_fs:.2f}", f"{usace_max_embedment_ft:.2f} ft"],
    })
    add_df_to_doc(doc, input_df)

    doc.add_paragraph("Table 1: Soil layer properties.")
    layer_report = layer_df.copy()
    layer_report.insert(0, "Top Depth (ft)", [0.0] + list(layer_report["Bottom Depth (ft)"].iloc[:-1]))
    add_df_to_doc(doc, layer_report)

    doc.add_heading("Theory and Calculation Method", level=2)
    theory_items = [
        "Vertical stress is calculated by summing the unit weight of each layer down to the calculation depth. Above the water table, moist unit weight is used; below the water table, saturated unit weight and hydrostatic pore pressure are used to compute effective vertical stress.",
        "Rankine active pressure is calculated as p'a = Ka sigma'v - 2c sqrt(Ka), with negative net active pressure clipped to zero.",
        "Coulomb active pressure uses wall inclination, wall friction, and backfill slope in the coefficient calculation.",
        "At-rest pressure uses a Jaky-type coefficient, K0 = (1 - sin(phi))(1 + sin(beta)).",
        "Passive effective soil pressure is calculated explicitly as p'p = sigma'v/Ka + 2c sqrt(1/Ka), beginning at the user-defined passive pressure start depth. The cohesion term is included only when the cohesion checkbox is selected.",
        "Passive total pressure is calculated as passive effective soil pressure plus pore-water pressure when total pressure is selected.",
        "Surcharge pressure is calculated separately from earth pressure. FHWA/WALLPRES strip loading, NAVFAC/Boussinesq point loading, and AASHTO 2:1 strip/isolated footing/point-load distribution are included as separate surcharge methods when selected.",
        "Elastic deflection is calculated from curvature M/EI. The moment per foot of wall is multiplied by the entered pile tributary width or spacing and converted to kip-in. Zero rotation and zero deflection are imposed at the selected virtual point of fixity. For the Shoring Suite-style option, the segment below virtual fixity is treated as restrained and its plotted rotation and deflection are set to zero. For the toe-fixed option, the same boundary is applied at the pile toe. Positive deflection is in the active-pressure direction.",
        "The optional USACE EM 1110-2-2504 cantilever method forms net-active and full-net-passive pressure distributions, reduces passive soil strengths by FSp, and solves simultaneous horizontal-force and moment equilibrium for the required penetration and transition point. The transition is used only as an idealized elastic fixity for the optional deflection diagram.",
    ]
    def add_safe_bullet(document, text):
        # Some company templates do not include Word's built-in "List Bullet" style.
        # Use it when available; otherwise write a normal paragraph with a bullet character.
        try:
            document.add_paragraph(text, style="List Bullet")
        except KeyError:
            document.add_paragraph(f"• {text}")

    for item in theory_items:
        add_safe_bullet(doc, item)

    doc.add_heading("Figures", level=2)
    figures = [
        ("Figure 1: Input geometry schematic.", create_geometry_figure()),
        ("Figure 2: Lateral earth pressure diagrams.", create_pressure_diagram_figure()),
        ("Figure 3: Active and at-rest pressure comparison.", create_active_comparison_figure()),
        ("Figure 4: Passive pressure comparison.", create_passive_comparison_figure()),
    ]
    next_fig_num = 5
    if surcharge_type != "None":
        figures.append((f"Figure {next_fig_num}: Separate surcharge pressure diagram.", create_surcharge_figure()))
        next_fig_num += 1
    figures.append((f"Figure {next_fig_num}: Net pressure, shear, moment, and elastic deflection diagrams along the pile/wall.", create_net_shear_moment_figure()))
    next_fig_num += 1
    if run_usace_analysis:
        figures.append((f"Figure {next_fig_num}: USACE cantilever equilibrium pressure, shear, moment, and approximate deflection diagrams.", create_usace_cantilever_figure()))
    for caption, fig in figures:
        doc.add_paragraph(caption)
        buf = fig_to_png_bytes(fig)
        doc.add_picture(buf, width=Inches(6.5))
        plt.close(fig)

    doc.add_heading("Results", level=2)
    doc.add_paragraph("Table 2: Summary of resultants.")
    add_df_to_doc(doc, summary_table_df())
    doc.add_paragraph("Table 3: Pressure values at selected depths.")
    add_df_to_doc(doc, report_table_df())
    if run_usace_analysis:
        doc.add_paragraph("Table 4: USACE cantilever equilibrium summary.")
        add_df_to_doc(doc, usace_summary_df())

    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(
        "The results are based on classical earth pressure methods and the input parameters entered by the user. The deflection calculation assumes a linear-elastic Euler-Bernoulli wall above an idealized virtual point of fixity; the wall below that point is treated as restrained for the Shoring Suite-style deflection plot. The USACE transition point is a rotational-stability construct near the zero-displacement point; using it as both zero rotation and zero deflection is an additional approximation for plotting elastic deflection. The USACE solver extends the last entered soil layer below the entered wall bottom when evaluating trial penetrations. This calculation does not iterate soil pressure with wall movement, include discrete braces or anchors, model nonlinear pile-soil springs, or account for cracking, yielding, shear deformation, construction staging, or second-order effects. The calculation is intended for preliminary engineering review and should be checked by the engineer of record before design use."
    )

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

# -----------------------------
# Results UI
# -----------------------------
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📐 Results & Diagrams", "📊 Comparison Charts", "📋 Formulas & Notes", "🔢 Detailed Tables", "🏛️ USACE Cantilever", "📝 Report"])

with tab1:
    st.subheader("Input Geometry Based on Current Parameters")
    fig_geo = create_geometry_figure()
    st.pyplot(fig_geo)
    plt.close(fig_geo)

    col_res, col_diag = st.columns([1, 1.6])
    with col_res:
        st.subheader("Layer Summary")
        display_layers = layer_df.copy()
        display_layers.insert(0, "Top Depth (ft)", [0.0] + list(display_layers["Bottom Depth (ft)"].iloc[:-1]))
        st.dataframe(display_layers, use_container_width=True, hide_index=True)

        st.subheader("Active Pressure Resultants")
        def result_card(method, cls, F, h):
            st.markdown(f"""
            <div class="result-card">
              <span class="method-badge {cls}">{method}</span>
              <h4>F<sub>a</sub> = {F:.2f} kips/ft</h4>
              <p style="margin:0;font-size:0.85rem;">Acts at <b>{h:.2f} ft</b> above base</p>
            </div>""", unsafe_allow_html=True)

        result_card("Rankine", "rankine", Fa_r, ha_r)
        result_card("Coulomb", "coulomb", Fa_c, ha_c)
        result_card("At-Rest", "atrest", Fa_0, ha_0)

        st.subheader("Surcharge Resultants - Separate")
        if surcharge_type == "None":
            st.info("No surcharge selected.")
        else:
            if surcharge_type in ["Uniform", "AASHTO"]:
                card_prefix = "Uniform surcharge" if surcharge_type == "Uniform" else "AASHTO live-load surcharge"
                result_card(f"{card_prefix} - Rankine K", "rankine", Fs_r, hs_r)
                result_card(f"{card_prefix} - Coulomb K", "coulomb", Fs_c, hs_c)
                result_card(f"{card_prefix} - At-Rest K", "atrest", Fs_0, hs_0)
            else:
                result_card(f"{surcharge_type} surcharge", "rankine", Fs_r, hs_r)
            st.caption("Combined active + surcharge totals are shown in the Detailed Tables tab.")

        st.subheader("Passive Pressure Resultants")
        st.markdown(f"""
        <div class="result-card"><span class="method-badge rankine">Rankine</span><h4>F<sub>p</sub> = {Fp_r:.2f} kips/ft</h4><p style="margin:0;font-size:0.85rem;">Acts at <b>{hp_r:.2f} ft</b> above base</p></div>
        <div class="result-card"><span class="method-badge coulomb">Coulomb</span><h4>F<sub>p</sub> = {Fp_c:.2f} kips/ft</h4><p style="margin:0;font-size:0.85rem;">Acts at <b>{hp_c:.2f} ft</b> above base</p></div>
        """, unsafe_allow_html=True)

        st.markdown(f"<div class='note'>Passive pressure is calculated only below <b>{passive_start_depth:.2f} ft</b> from the top.</div>", unsafe_allow_html=True)
        if include_water:
            st.markdown(f"<div class='note'>Water pressure is added below <b>{water_table:.2f} ft</b> when total pressure is shown.</div>", unsafe_allow_html=True)

        st.subheader("Elastic Deflection Results")
        st.markdown(f"""
        <div class="result-card"><span class="method-badge rankine">Rankine</span><h4>Top y = {top_deflection_rankine:.3f} in</h4><p style="margin:0;font-size:0.85rem;">Maximum |y| = <b>{max_abs_deflection_rankine:.3f} in</b> at depth {max_deflection_depth_rankine:.2f} ft</p></div>
        <div class="result-card"><span class="method-badge coulomb">Coulomb</span><h4>Top y = {top_deflection_coulomb:.3f} in</h4><p style="margin:0;font-size:0.85rem;">Maximum |y| = <b>{max_abs_deflection_coulomb:.3f} in</b> at depth {max_deflection_depth_coulomb:.2f} ft</p></div>
        """, unsafe_allow_html=True)
        st.caption(
            f"{deflection_boundary_label}; E = {youngs_modulus_ksi:,.1f} ksi, "
            f"I = {moment_of_inertia_in4:,.2f} in⁴, tributary width/spacing = "
            f"{pile_tributary_width_ft:.2f} ft."
        )

    with col_diag:
        fig, axes = plt.subplots(1, 3, figsize=(12, 7), sharey=True)
        fig.patch.set_facecolor("#f8fafc")
        configs = [
            ("Rankine", pa_rankine, pp_rankine, "#3b82f6", "#f59e0b"),
            ("Coulomb", pa_coulomb, pp_coulomb, "#22c55e", "#ef4444"),
            ("At-Rest", pa_atrest, pa_atrest, "#a855f7", "#a855f7"),
        ]
        for ax, (title, pa, pp, c_a, c_p) in zip(axes, configs):
            ax.set_facecolor("#f8fafc")
            ax.fill_betweenx(depths, 0, pp, alpha=0.15, color=c_p)
            ax.plot(pp, depths, color=c_p, lw=2, label="Passive" if title != "At-Rest" else "At-Rest")
            ax.fill_betweenx(depths, 0, -pa, alpha=0.20, color=c_a)
            ax.plot(-pa, depths, color=c_a, lw=2.2, label="Active")
            for b in layer_df["Bottom Depth (ft)"].iloc[:-1]:
                ax.axhline(b, color="#64748b", lw=0.8, ls="--", alpha=0.6)
            if passive_start_depth > 0:
                ax.axhline(passive_start_depth, color="#92400e", lw=1.2, ls="-.", alpha=0.9)
            if include_water:
                ax.axhline(water_table, color="#0ea5e9", lw=1.2, ls=":", alpha=0.9)
            ax.axvline(0, color="#334155", lw=0.8)
            ax.set_title(title, fontweight="bold", fontsize=12, color="#1a3a5c")
            ax.set_xlabel("Pressure (psf)", fontsize=9)
            ax.legend(fontsize=8, loc="lower right")
            ax.grid(True, alpha=0.3, linestyle="--")
            ax.tick_params(labelsize=8)
        # Shared depth axis: zero at top, depth increasing downward.
        axes[0].set_ylim(H, 0)
        axes[0].set_ylabel("Depth (ft)", fontsize=10)
        fig.suptitle("Lateral Earth Pressure Diagrams", fontsize=13, fontweight="bold", color="#1a3a5c", y=1.01)
        plt.tight_layout()
        st.pyplot(fig)
        plt.close(fig)

with tab2:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Active / At-Rest Pressure")
        fig3, ax3 = plt.subplots(figsize=(6, 6))
        ax3.plot(pa_rankine, depths, lw=2.2, label="Rankine Active")
        ax3.plot(pa_coulomb, depths, lw=2.2, ls="--", label="Coulomb Active")
        ax3.plot(pa_atrest, depths, lw=2.2, ls=":", label="At-Rest")
        ax3.set_ylim(H, 0); ax3.set_xlabel("Pressure (psf)"); ax3.set_ylabel("Depth (ft)")
        ax3.legend(); ax3.grid(True, alpha=0.3); ax3.set_title("Lateral Pressure Comparison")
        fig3.tight_layout(); st.pyplot(fig3); plt.close(fig3)
    with col2:
        st.subheader("Passive Pressure")
        fig4, ax4 = plt.subplots(figsize=(6, 6))
        ax4.plot(pp_rankine, depths, lw=2.2, label="Rankine Passive")
        ax4.plot(pp_coulomb, depths, lw=2.2, ls="--", label="Coulomb Passive")
        ax4.set_ylim(H, 0); ax4.set_xlabel("Pressure (psf)"); ax4.set_ylabel("Depth (ft)")
        ax4.legend(); ax4.grid(True, alpha=0.3); ax4.set_title("Passive Pressure Comparison")
        fig4.tight_layout(); st.pyplot(fig4); plt.close(fig4)

    st.subheader("Net Pressure, Shear, Moment, and Deflection Along Pile/Wall")
    fig_sm = create_net_shear_moment_figure()
    st.pyplot(fig_sm)
    plt.close(fig_sm)
    st.caption(
        "Net pressure = active pressure + surcharge pressure - passive pressure. Shear is integrated from the free top downward, "
        "moment is the integral of shear, and deflection is obtained by double integration of M/EI above the selected virtual fixity point. "
        "Pressure, shear, and moment are reported per foot of wall; pile deflection uses the entered tributary width/spacing. "
        f"The deflection boundary condition is: {deflection_boundary_label}. Positive deflection is in the active-pressure direction; "
        "negative deflection is toward the passive side."
    )

    csm1, csm2 = st.columns(2)
    with csm1:
        st.metric("Max |V| - Rankine", f"{max_abs_shear_rankine:.2f} kips/ft")
        st.metric("Max |M| - Rankine", f"{max_abs_moment_rankine:.2f} kip-ft/ft")
        st.metric("Max |y| - Rankine", f"{max_abs_deflection_rankine:.3f} in")
    with csm2:
        st.metric("Max |V| - Coulomb", f"{max_abs_shear_coulomb:.2f} kips/ft")
        st.metric("Max |M| - Coulomb", f"{max_abs_moment_coulomb:.2f} kip-ft/ft")
        st.metric("Max |y| - Coulomb", f"{max_abs_deflection_coulomb:.3f} in")

    if surcharge_type != "None":
        st.subheader("Surcharge Pressure - Separate")
        fig5, ax5 = plt.subplots(figsize=(6, 6))
        if surcharge_type == "Uniform":
            ax5.plot(sur_r, depths, lw=2.2, label="Rankine K × q")
            ax5.plot(sur_c, depths, lw=2.2, ls="--", label="Coulomb K × q")
            ax5.plot(sur_0, depths, lw=2.2, ls=":", label="At-Rest K × q")
        elif surcharge_type == "AASHTO":
            ax5.plot(sur_r, depths, lw=2.2, label="Rankine K × Δσv")
            ax5.plot(sur_c, depths, lw=2.2, ls="--", label="Coulomb K × Δσv")
            ax5.plot(sur_0, depths, lw=2.2, ls=":", label="At-Rest K × Δσv")
        else:
            ax5.plot(sur_r, depths, lw=2.2, label=surcharge_type)
        ax5.set_ylim(H, 0); ax5.set_xlabel("Surcharge Pressure (psf)"); ax5.set_ylabel("Depth (ft)")
        ax5.legend(); ax5.grid(True, alpha=0.3); ax5.set_title("Surcharge Pressure on Wall")
        fig5.tight_layout(); st.pyplot(fig5); plt.close(fig5)

    st.subheader("Resultant Force Comparison")
    fig6, axes6 = plt.subplots(1, 2, figsize=(10, 4))
    methods = ["Rankine", "Coulomb", "At-Rest"]
    bars = axes6[0].bar(methods, [Fa_r, Fa_c, Fa_0])
    axes6[0].bar_label(bars, fmt="%.2f", padding=3, fontsize=9)
    axes6[0].set_title("Active / At-Rest Resultant (kips/ft)"); axes6[0].set_ylabel("kips/ft"); axes6[0].grid(axis="y", alpha=0.3)
    bars2 = axes6[1].bar(methods[:2], [Fp_r, Fp_c])
    axes6[1].bar_label(bars2, fmt="%.2f", padding=3, fontsize=9)
    axes6[1].set_title("Passive Resultant (kips/ft)"); axes6[1].set_ylabel("kips/ft"); axes6[1].grid(axis="y", alpha=0.3)
    fig6.tight_layout(); st.pyplot(fig6); plt.close(fig6)

with tab3:
    st.markdown(r"""
### How this version handles layers and groundwater

- The app computes vertical stress by integrating each soil layer from the surface to the selected depth.
- Above the water table it uses moist unit weight, γm.
- Below the water table it uses saturated unit weight, γsat, computes pore pressure, then uses effective vertical stress for soil pressure.
- Total lateral pressure = effective lateral soil pressure + water pressure.
- Passive pressure can start at any selected depth from the wall top; pressure above that depth is set to zero.
- Passive effective soil pressure is calculated as $p'_p = \\sigma'_{v,p}/K_a + 2c\\sqrt{1/K_a}$ when cohesion is enabled. For $c=0$, this is exactly $\\sigma'_{v,p}/K_a$.
- Passive total pressure is the passive effective soil pressure plus passive-side pore-water pressure when total pressure is selected.
- Surcharge pressure is calculated separately and is not mixed into the basic earth-pressure resultants.
- For strip surcharge, the app follows the FHWA/WALLPRES-style rigid-wall equation: Δp = 2q/π × [β - sin(β)cos(2α)], where β = atan(x2/z) - atan(x1/z), α = atan(x1/z) + β/2, and x2 = x1 + strip width.
- For AASHTO surcharge, the app sets surcharge pressure to zero above z2 and calculates it only below z2. Below z2, z2 = 2d - bf and D1 = (bf + z)/2 + d. For strip footing, Δσv = Pv / D1. For isolated rectangular footing, Δσv = P'v / [D1(L+z)]. For point load, bf = 0 and Δσv = P'v / D1². It calculates Δσv separately and then applies Δp = K × Δσv.
- If cohesion creates negative active pressure, the active pressure is clipped to zero for the net diagram.
- Deflection uses zero rotation and zero displacement at the selected virtual point of fixity. Curvature is integrated only above that point; the embedded segment below it is treated as restrained and plotted at zero deflection. In the Shoring Suite-style option, the fixity point is entered as a percentage of embedment below the excavation/passive-pressure start depth; the default is 60%.

### English-unit conventions

| Quantity | Unit |
|---|---|
| Depth / wall height | ft |
| Soil unit weight | pcf |
| Cohesion | psf |
| Pressure | psf |
| Uniform surcharge | psf |
| Strip surcharge | psf over a strip width in ft |
| Line load | lb/ft |
| Point load | lb |
| AASHTO strip load Pv | lb/ft |
| AASHTO isolated footing P'v | lb |
| Resultant force | kips/ft |
| Young's modulus E | ksi |
| Moment of inertia I | in⁴ |
| Rotation | radians |
| Deflection | in |

### Key formulas

Rankine active pressure, effective-stress form:

$$p'_a = K_a \sigma'_v - 2c\sqrt{K_a}$$

Rankine passive pressure, effective-stress form:

$$p'_p = \frac{\sigma'_{v,p}}{K_a} + 2c\sqrt{\frac{1}{K_a}}$$

For cohesionless soil:

$$p'_p = \frac{\sigma'_{v,p}}{K_a}$$

Total pressure below groundwater:

$$p_{total} = p' + u$$

Hydrostatic pore pressure:

$$u = \gamma_w(z-z_w)$$

where $z_w$ is the water table depth.

Elastic beam curvature and deflection:

$$\kappa = \frac{M}{EI}, \qquad \theta = \int \kappa\,dx, \qquad y = \int \theta\,dx$$

For the point-of-fixity option, the boundary conditions are $\theta(z_f)=0$ and $y(z_f)=0$, where $z_f$ is measured from the wall top. Curvature is integrated upward from $z_f$ to the wall top. The embedded segment below $z_f$ is treated as restrained, so its plotted rotation and deflection are zero. The point is selected as a user-entered fraction of the embedment below the excavation/passive-pressure start depth. For the toe-fixed option, $z_f=H$. The moment used for one pile is the moment per foot of wall multiplied by the entered pile tributary width or spacing. Positive deflection is in the active-pressure direction; negative deflection is toward the passive side.

### USACE cantilever-wall equilibrium method

The optional method follows USACE EM 1110-2-2504, Chapter 5. Passive strengths are reduced as:

$$\tan\phi_{eff}=\frac{\tan\phi}{FS_p}, \qquad c_{eff}=\frac{c}{FS_p}$$

The design pressure equals the net-active distribution from the top to the transition point. From the transition to the design bottom, pressure varies linearly to the full net-passive pressure. Required penetration $d$ and the distance $z$ from the design bottom to the transition are solved from:

$$\sum F_H=0, \qquad \sum M=0$$

The manual describes the transition as near the point of zero displacement. In this app, using the transition as both zero rotation and zero deflection for the optional elastic-deflection plot is an additional idealization, not part of the USACE equilibrium equations.
""")

with tab4:
    st.subheader("Pressure Values at Selected Depths")
    idx = np.linspace(0, len(depths) - 1, min(30, len(depths)), dtype=int)
    df = pd.DataFrame({
        "Depth (ft)": np.round(depths[idx], 2),
        "Layer": [layer_index_at_depth(float(z)) + 1 for z in depths[idx]],
        "σv total (psf)": np.round(sigma_v_total[idx], 1),
        "u water (psf)": np.round(u_water[idx], 1),
        "σv effective (psf)": np.round(sigma_v_eff[idx], 1),
        "Passive σv effective from start (psf)": np.round(sigma_v_eff_passive[idx], 1),
        "Ka Rankine": np.round(Ka_r[idx], 3),
        "Ka Coulomb": np.round(Ka_c[idx], 3),
        "Rankine active total (psf)": np.round(pa_rankine[idx], 1),
        "Coulomb active total (psf)": np.round(pa_coulomb[idx], 1),
        "At-rest total (psf)": np.round(pa_atrest[idx], 1),
        "Surcharge Rankine/rigid (psf)": np.round(sur_r[idx], 1),
        "Surcharge Coulomb (psf)": np.round(sur_c[idx], 1),
        "Rankine active + surcharge (psf)": np.round(pa_rankine_plus_surcharge[idx], 1),
        "Coulomb active + surcharge (psf)": np.round(pa_coulomb_plus_surcharge[idx], 1),
        "At-rest + surcharge (psf)": np.round(pa_atrest_plus_surcharge[idx], 1),
        "Rankine passive effective (psf)": np.round(pp_rankine_eff[idx], 1),
        "Coulomb passive effective (psf)": np.round(pp_coulomb_eff[idx], 1),
        "Rankine passive total (psf)": np.round(pp_rankine[idx], 1),
        "Coulomb passive total (psf)": np.round(pp_coulomb[idx], 1),
        "Rankine net pressure (psf)": np.round(net_pressure_rankine[idx], 1),
        "Coulomb net pressure (psf)": np.round(net_pressure_coulomb[idx], 1),
        "Rankine shear (kips/ft)": np.round(shear_rankine[idx], 3),
        "Coulomb shear (kips/ft)": np.round(shear_coulomb[idx], 3),
        "Rankine moment (kip-ft/ft)": np.round(moment_rankine[idx], 3),
        "Coulomb moment (kip-ft/ft)": np.round(moment_coulomb[idx], 3),
        "Rankine rotation (rad)": np.round(rotation_rankine[idx], 6),
        "Coulomb rotation (rad)": np.round(rotation_coulomb[idx], 6),
        "Rankine deflection (in)": np.round(deflection_rankine[idx], 4),
        "Coulomb deflection (in)": np.round(deflection_coulomb[idx], 4),
    })
    st.dataframe(df, use_container_width=True, hide_index=True)

    st.subheader("Summary of Results")
    summary = pd.DataFrame({
        "Method": ["Rankine", "Coulomb", "At-Rest"],
        "Earth/Water Fa or F0 (kips/ft)": [f"{Fa_r:.2f}", f"{Fa_c:.2f}", f"{Fa_0:.2f}"],
        "Earth/Water height above base (ft)": [f"{ha_r:.2f}", f"{ha_c:.2f}", f"{ha_0:.2f}"],
        "Separate surcharge F (kips/ft)": [f"{Fs_r:.2f}", f"{Fs_c:.2f}", f"{Fs_0:.2f}"],
        "Surcharge height above base (ft)": [f"{hs_r:.2f}", f"{hs_c:.2f}", f"{hs_0:.2f}"],
        "Combined active + surcharge F (kips/ft)": [f"{Ft_r:.2f}", f"{Ft_c:.2f}", f"{Ft_0:.2f}"],
        "Combined height above base (ft)": [f"{ht_r:.2f}", f"{ht_c:.2f}", f"{ht_0:.2f}"],
        "Fp (kips/ft)": [f"{Fp_r:.2f}", f"{Fp_c:.2f}", "—"],
        "Passive height above base (ft)": [f"{hp_r:.2f}", f"{hp_c:.2f}", "—"],
        "Max |V| (kips/ft)": [f"{max_abs_shear_rankine:.2f}", f"{max_abs_shear_coulomb:.2f}", "—"],
        "Max |M| (kip-ft/ft)": [f"{max_abs_moment_rankine:.2f}", f"{max_abs_moment_coulomb:.2f}", "—"],
        "Top deflection (in)": [f"{top_deflection_rankine:.3f}", f"{top_deflection_coulomb:.3f}", "—"],
        "Max |deflection| (in)": [f"{max_abs_deflection_rankine:.3f}", f"{max_abs_deflection_coulomb:.3f}", "—"],
        "Depth of max |deflection| (ft)": [f"{max_deflection_depth_rankine:.2f}", f"{max_deflection_depth_coulomb:.2f}", "—"],
        "Point of fixity depth (ft)": [f"{point_of_fixity_depth:.2f}", f"{point_of_fixity_depth:.2f}", "—"],
    })
    st.dataframe(summary, use_container_width=True, hide_index=True)

    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button("⬇️ Download Pressure Table (CSV)", df.to_csv(index=False), "pressure_data_english.csv", "text/csv")
    with col_dl2:
        st.download_button("⬇️ Download Summary Table (CSV)", summary.to_csv(index=False), "summary_results_english.csv", "text/csv")


with tab5:
    st.subheader("USACE EM 1110-2-2504 Cantilever-Wall Method")
    st.markdown(
        "The method uses the net-active pressure distribution from the wall top to a transition point. "
        "Below the transition, the design pressure varies linearly to the full net-passive pressure at "
        "the design bottom. The solver varies penetration and transition location until both horizontal-force "
        "and moment equilibrium are satisfied. Passive coefficients remain consistent with the app's "
        "Kp = 1/Ka convention after strength reduction."
    )
    st.caption(
        f"Passive strength reduction: FSp = {usace_passive_fs:.2f}; "
        f"maximum trial embedment = {usace_max_embedment_ft:.2f} ft. "
        "The last entered soil layer is extended for trial depths below the entered wall bottom."
    )

    usace_summary = usace_summary_df()
    st.dataframe(usace_summary, use_container_width=True, hide_index=True)

    for method, solution, css_class in [
        ("Rankine", usace_rankine, "rankine"),
        ("Coulomb", usace_coulomb, "coulomb"),
    ]:
        if solution.get("success", False):
            entered_embedment = max(float(H) - float(passive_start_depth), 0.0)
            adequacy_text = (
                "Entered wall penetration is adequate for this calculated requirement."
                if entered_embedment + 1.0e-6 >= solution["embedment"]
                else (
                    f"Entered penetration is {solution['embedment'] - entered_embedment:.2f} ft "
                    "less than the calculated USACE requirement."
                )
            )
            st.markdown(
                f"""
                <div class="result-card">
                  <span class="method-badge {css_class}">{method} USACE</span>
                  <h4>Required penetration d = {solution['embedment']:.2f} ft</h4>
                  <p style="margin:0;font-size:0.85rem;">
                    Design bottom = <b>{solution['bottom_depth']:.2f} ft</b> from top ·
                    Transition = <b>{solution['transition_depth']:.2f} ft</b> from top ·
                    z above bottom = <b>{solution['z_above_bottom']:.2f} ft</b><br>
                    {adequacy_text}
                  </p>
                </div>
                """,
                unsafe_allow_html=True,
            )
        else:
            st.warning(f"{method}: {solution.get('message', 'USACE solution unavailable.')}")

    if run_usace_analysis:
        fig_usace = create_usace_cantilever_figure()
        st.pyplot(fig_usace)
        plt.close(fig_usace)
        st.caption(
            "For the USACE pressure diagram, the transition point is solved from force and moment equilibrium. "
            "For the deflection panel only, that transition is idealized as a point of zero rotation and zero "
            "deflection so M/EI can be integrated; EM 1110-2-2504 describes the transition as near the point of "
            "zero displacement and does not by itself define a complete elastic-deformation boundary condition."
        )

    col_usace_1, col_usace_2, col_usace_3 = st.columns(3)
    with col_usace_1:
        st.download_button(
            "⬇️ Download USACE Summary (CSV)",
            usace_summary.to_csv(index=False),
            "usace_cantilever_summary.csv",
            "text/csv",
            use_container_width=True,
        )
    with col_usace_2:
        rankine_usace_df = usace_detail_df(usace_rankine)
        if not rankine_usace_df.empty:
            st.download_button(
                "⬇️ Rankine USACE Data (CSV)",
                rankine_usace_df.to_csv(index=False),
                "usace_rankine_response.csv",
                "text/csv",
                use_container_width=True,
            )
    with col_usace_3:
        coulomb_usace_df = usace_detail_df(usace_coulomb)
        if not coulomb_usace_df.empty:
            st.download_button(
                "⬇️ Coulomb USACE Data (CSV)",
                coulomb_usace_df.to_csv(index=False),
                "usace_coulomb_response.csv",
                "text/csv",
                use_container_width=True,
            )


with tab6:
    st.subheader("Generate Technical Memorandum Report")
    st.caption("The report uses the TSG technical memo format when a template is uploaded or when 'TSG Technical MEMO 2026.docx' is available beside the app file.")

    col_meta1, col_meta2 = st.columns(2)
    with col_meta1:
        report_date = st.date_input("Report Date", value=date.today())
        report_to = st.text_input("To", value="")
        report_cc = st.text_input("Cc", value="")
    with col_meta2:
        report_from = st.text_input("From", value="")
        report_project = st.text_input("Project", value="Project")
        report_subject = st.text_input("Subject", value="Lateral Earth Pressure Calculation")

    uploaded_template = st.file_uploader("Optional Word template (.docx)", type=["docx"])

    if st.button("📝 Generate Word Report", use_container_width=True):
        try:
            template_bytes = uploaded_template.getvalue() if uploaded_template is not None else None
            report_bytes = generate_word_report(
                template_bytes=template_bytes,
                meta={
                    "date": report_date.isoformat(),
                    "to": report_to,
                    "cc": report_cc,
                    "from": report_from,
                    "project": report_project,
                    "subject": report_subject,
                },
            )
            st.session_state["generated_report_bytes"] = report_bytes
            st.success("Report generated successfully.")
        except Exception as exc:
            st.error(f"Could not generate report: {exc}")

    if "generated_report_bytes" in st.session_state:
        st.download_button(
            "⬇️ Download Word Report",
            data=st.session_state["generated_report_bytes"],
            file_name="lateral_earth_pressure_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

st.markdown("---")
st.markdown("""
<div style='text-align:center;color:#64748b;font-size:0.85rem;padding:10px'>
  🏗️ Lateral Earth Pressure and Elastic Deflection Calculator | English Units | Multi-Layer Soil | Groundwater | Rankine · Coulomb · At-Rest · Separate Surcharge · USACE Cantilever
</div>
""", unsafe_allow_html=True)
